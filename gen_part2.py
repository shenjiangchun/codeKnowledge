"""Part 2: Spring Boot + PostgreSQL + Redis + Kafka + MyBatis"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = r'C:\Users\47583\projects\hisi_dev_tool v5.0\沈江春_面试准备手册_完整版.docx'
doc = Document(OUTPUT)

def P(text, bold=False, sz=None, clr=None, align=None, sa=None, sn=None):
    p = doc.add_paragraph(style=sn) if sn else doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if sz: r.font.size = Pt(sz)
    if clr: r.font.color.rgb = RGBColor.from_string(clr)
    if align: p.alignment = align
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    return p
def B(text, lv=0):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    r = p.add_run(text); r.font.size = Pt(10)
    if lv: p.paragraph_format.left_indent = Cm(1.5*lv)
def TBL(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        c.paragraphs[0].add_run(h).bold = True; c.paragraphs[0].runs[0].font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            c.paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()
def CODE(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
def A(title): P(title, bold=True, sz=11, clr='2E75B6', sa=2)
def ANS(text): P(text, sz=10, sa=6)
def ABL(text): B(text)
def Q(text): P(text, bold=True, sz=11, sa=4)
def SEC1(t): doc.add_heading(t, level=1)
def SEC2(t): doc.add_heading(t, level=2)
def SEC3(t): doc.add_heading(t, level=3)

# ==================== 二、Spring Boot ====================
SEC2('二、Spring Boot 框架')
SEC3('2.1 IoC / DI / AOP')

Q('Q：Spring IoC 的原理？Bean 的生命周期？')
P('Bean 生命周期（完整 8 步）：', bold=True, sz=10)
B('1. 实例化（Instantiation）：容器通过反射创建 Bean 实例（调用构造器）。')
B('2. 属性注入（Populate Properties）：@Autowired / @Value / @Resource 注入依赖。')
B('3. Aware 回调：如果 Bean 实现了 BeanNameAware、BeanFactoryAware、ApplicationContextAware 等接口，容器依次回调 setBeanName()、setBeanFactory()、setApplicationContext()。')
B('4. BeanPostProcessor.postProcessBeforeInitialization：容器级别的前置处理。@PostConstruct 注解的方法在此阶段执行（CommonAnnotationBeanPostProcessor）。')
B('5. InitializingBean.afterPropertiesSet()：如果 Bean 实现了 InitializingBean 接口，调用 afterPropertiesSet()。或者通过 @Bean(initMethod=\"init\") 指定初始化方法。')
B('6. BeanPostProcessor.postProcessAfterInitialization：容器级别的后置处理。AOP 代理对象在此阶段生成（AbstractAutoProxyCreator）。')
B('7. Bean 就绪，可以被使用。')
B('8. 销毁：容器关闭时，调用 @PreDestroy -> DisposableBean.destroy() -> @Bean(destroyMethod=\"destroy\")。')

A('追问 1：@Autowired 的注入方式？推荐哪种？')
ANS('Spring 支持三种注入方式：')
ABL('构造器注入（推荐）：@Autowired 加在构造器上。优点：(1) 依赖不可变（可以声明为 final）；(2) 必须在构造时提供所有依赖，不会出现 NPE；(3) 方便单元测试（直接 new 传入 mock）。Spring 官方推荐，Spring Boot 4.x 可能强制要求。')
ABL('字段注入（不推荐）：@Autowired 加在字段上。缺点：(1) 依赖可变（非 final）；(2) 隐藏了依赖关系（看构造器不知道需要哪些依赖）；(3) 难以单元测试（必须用反射注入 mock）；(4) 可能导致循环依赖的隐蔽问题。')
ABL('Setter 注入（可选）：@Autowired 加在 setter 方法上。适用于可选依赖（可以不调用 setter）。')
ANS('Spring Boot 3 中如果只有一个构造器，可以省略 @Autowired 注解（隐式注入）。')

A('追问 2：循环依赖如何解决？')
ANS('Spring 通过三级缓存解决单例 Bean 的 setter/字段注入循环依赖：')
ABL('一级缓存 singletonObjects：完整的 Bean 实例（已初始化、已注入属性）。')
ABL('二级缓存 earlySingletonObjects：提前暴露的半成品 Bean（已实例化但未注入属性），用于解决循环依赖。')
ABL('三级缓存 singletonFactories：Bean 的 ObjectFactory（lambda 表达式），在需要时生成早期引用或代理对象。')
ANS('解决流程（A 依赖 B，B 依赖 A）：(1) 创建 A：实例化 A 后，将 A 的 ObjectFactory 放入三级缓存；(2) 注入 A 的属性时发现依赖 B；(3) 创建 B：实例化 B 后放入三级缓存；(4) 注入 B 的属性时发现依赖 A，从三级缓存获取 A 的 ObjectFactory，调用 getObject() 得到 A 的早期引用（如果 A 需要 AOP 代理则在此时生成代理），放入二级缓存，注入给 B；(5) B 创建完成，放入一级缓存；(6) 回到 A，注入 B，A 创建完成。')
ANS('Spring Boot 3 中默认不允许循环依赖（spring.main.allow-circular-references=false）。如果遇到循环依赖，应该优先通过重构代码解决（提取公共服务、使用 @Lazy 延迟注入、使用事件驱动解耦），而不是开启循环依赖。构造器注入的循环依赖无法通过三级缓存解决（因为对象还没创建完就需要依赖），会直接报错。')

A('追问 3：Spring Boot 3 中循环依赖还支持吗？')
ANS('Spring Boot 3 默认禁止循环依赖（spring.main.allow-circular-references=false）。这是 Spring 团队有意为之的设计决策：循环依赖通常是设计问题的信号。')
ANS('如果确实需要处理，有三种方案：')
ABL('方案一：@Lazy 注解。在注入点加 @Lazy，注入一个延迟代理，实际 Bean 在首次使用时才初始化。@Autowired @Lazy private B b;')
ABL('方案二：使用 @EventListener / ApplicationEventPublisher 事件驱动解耦。A 发布事件，B 监听事件，两者不再直接依赖。')
ABL('方案三：提取公共服务 C，A 和 B 都依赖 C，消除循环。')

SEC3('2.2 Spring Boot 自动配置')

Q('Q：@SpringBootApplication 的组成？自动配置原理？')
CODE('@SpringBootApplication\n= @SpringBootConfiguration  // 标记为配置类\n+ @EnableAutoConfiguration  // 开启自动配置\n+ @ComponentScan            // 扫描当前包及子包')
ANS('自动配置原理（Spring Boot 3.x）：')
B('1. @EnableAutoConfiguration 通过 @Import(AutoConfigurationImportSelector.class) 导入选择器。')
B('2. AutoConfigurationImportSelector 读取 META-INF/spring/org.springframework.boot.autoconfigure.AutoConfiguration.imports 文件（Boot 2.x 是 META-INF/spring.factories）。')
B('3. 文件中列出了所有自动配置类（如 DataSourceAutoConfiguration、RedisAutoConfiguration 等），共 100+ 个。')
B('4. 每个自动配置类通过 @Conditional 系列注解按需加载：@ConditionalOnClass（classpath 中存在某个类）、@ConditionalOnBean（容器中存在某个 Bean）、@ConditionalOnMissingBean（容器中不存在某个 Bean）、@ConditionalOnProperty（配置项为某个值）。')
B('5. 例如 DataSourceAutoConfiguration：如果 classpath 中有 javax.sql.DataSource 类且用户没有自定义 DataSource Bean，则自动配置 HikariCP 连接池。')

A('追问 1：如何自定义一个 Starter？')
ANS('自定义 Starter 的步骤（以 my-service-spring-boot-starter 为例）：')
B('1. 创建 autoconfigure 模块（my-service-spring-boot-autoconfigure）：编写配置类 @Configuration + @ConditionalOnXxx、属性类 @ConfigurationProperties(prefix=\"my.service\")、核心服务类。')
B('2. 在 autoconfigure 模块的 META-INF/spring/org.springframework.boot.autoconfigure.AutoConfiguration.imports 文件中注册配置类全限定名。')
B('3. 创建 starter 模块（my-service-spring-boot-starter）：只包含一个 pom.xml，依赖 autoconfigure 模块。这是约定：starter 只做依赖聚合，不写代码。')
B('4. 发布到 Maven 仓库，其他项目只需引入 starter 依赖即可自动生效。')

SEC3('2.3 Spring Boot 3 / Spring 6 新特性')

Q('Q：Spring Boot 3 的主要变化？')
B('1. Jakarta EE 9+：javax.* -> jakarta.*，这是最大的迁移成本。所有 Servlet、JPA、Bean Validation 等 API 的包名都变了。')
B('2. Java 17 最低要求：不能在 JDK 8/11 上运行。')
B('3. AOT（Ahead-of-Time）编译：支持 GraalVM Native Image，应用编译为原生二进制文件，启动时间从秒级降到毫秒级，内存占用大幅降低。')
B('4. Observability（可观测性）：内置 Micrometer + OpenTelemetry 集成，支持 Metrics/Tracing/Logging 三大支柱的标准化采集。')
B('5. Problem Detail（RFC 7807）：标准化的 REST API 错误响应格式，包含 type/title/status/detail/instance 字段。')
B('6. HTTP Interface Client：类似 Feign 的声明式 HTTP 客户端，用接口 + 注解定义远程调用。@GetExchange(\"/users/{id}\") User getUser(@PathVariable Long id);')
B('7. @ConfigurationProperties 支持 record 类型。')
B('8. Spring Security 配置简化：WebSecurityConfigurerAdapter 被移除，改用 SecurityFilterChain Bean 方式。')

A('追问 1：javax -> jakarta 迁移策略？')
ANS('迁移步骤：')
B('1. 全局替换：javax.servlet -> jakarta.servlet, javax.persistence -> jakarta.persistence, javax.validation -> jakarta.validation 等。可以用 IDE 批量替换或 OpenRewrite 自动化工具。')
B('2. 三方库升级：Hibernate 5.x -> 6.x, Tomcat 9 -> 10+, Jackson 2.14+（兼容 jakarta）。检查每个三方库的 Jakarta 兼容版本。')
B('3. 编译期检查：替换后编译，所有遗漏的 javax 引用都会报错。')
B('4. 运行时检查：某些框架（如 Spring）在运行时扫描注解，编译通过不代表运行时正常，需要全面测试。')
ANS('OpenRewrite 自动化迁移：引入 org.openrewrite.recipe:rewrite-spring-boot-recipes 依赖，运行 mvn rewrite:run 或 gradle rewriteRun，自动完成大部分迁移工作。')

SEC3('2.4 事务管理')

Q('Q：Spring 事务传播行为有哪些？')
TBL(['传播行为','说明','典型场景'],[
    ['REQUIRED（默认）','有事务则加入，无则新建','大多数业务方法'],
    ['REQUIRES_NEW','无论如何新建事务，挂起当前事务','独立的子操作（如日志记录、审计），不受外层回滚影响'],
    ['NESTED','嵌套事务（savepoint）','部分回滚：内层回滚不影响外层'],
    ['SUPPORTS','有则加入，无则非事务执行','查询方法'],
    ['NOT_SUPPORTED','非事务执行，挂起当前事务','不需要事务的操作'],
    ['MANDATORY','必须在事务中，否则抛异常','强制要求调用方开事务'],
    ['NEVER','必须非事务，否则抛异常','强制要求调用方不开事务'],
])

A('追问 1：@Transactional 在什么情况下会失效？')
ANS('Spring 事务基于 AOP 代理实现，以下情况会导致失效：')
ABL('自调用（最常见）：同一个类中 A 方法调用 B 方法，B 方法加了 @Transactional，但不会生效。因为自调用走的是 this.method()，不经过代理对象。解决：注入自身代理（@Lazy self）、拆分到不同类、使用 AopContext.currentProxy()。')
ABL('非 public 方法：Spring AOP 默认只代理 public 方法。private/protected/package-private 方法上的 @Transactional 会被忽略（不报错但不生效）。')
ABL('异常被 catch 吞掉：如果方法内部 catch 了异常没有重新抛出，Spring 无法感知到异常，不会回滚。解决：catch 块中 throw new RuntimeException(e) 或指定 rollbackFor。')
ABL('异常类型不匹配：默认只在 RuntimeException 和 Error 时回滚，Checked Exception 不回滚。解决：@Transactional(rollbackFor = Exception.class)。')
ABL('数据库不支持事务：MySQL 的 MyISAM 引擎不支持事务；DDL 操作（CREATE/ALTER/DROP）在某些数据库中会隐式提交事务。')
ABL('Bean 未被 Spring 管理：手动 new 的对象不走 AOP 代理。')
ABL('多线程场景：新线程中的操作不在同一个事务中（事务信息存储在 ThreadLocal 中，子线程无法获取父线程的事务上下文）。')

A('追问 2：事务隔离级别？脏读、不可重复读、幻读？')
ANS('SQL 标准定义了四种隔离级别，从低到高：')
TBL(['隔离级别','脏读','不可重复读','幻读','说明'],[
    ['READ_UNCOMMITTED','可能','可能','可能','最低级别，几乎不用'],
    ['READ_COMMITTED','不可能','可能','可能','Oracle/PostgreSQL 默认'],
    ['REPEATABLE_READ','不可能','不可能','可能（InnoDB 用 MVCC+Gap Lock 解决）','MySQL InnoDB 默认'],
    ['SERIALIZABLE','不可能','不可能','不可能','最严格，性能最差'],
])
ANS('三种问题的定义：')
ABL('脏读（Dirty Read）：事务 A 读取了事务 B 尚未提交的数据，B 随后回滚，A 读到了不存在的数据。')
ABL('不可重复读（Non-Repeatable Read）：事务 A 第一次读取某行数据，事务 B 修改了该行并提交，事务 A 再次读取得到不同结果。重点在于「已有数据被修改」。')
ABL('幻读（Phantom Read）：事务 A 按条件查询得到 N 行，事务 B 插入了满足条件的新行并提交，事务 A 再次查询得到 N+1 行。重点在于「新增了行」。')
ANS('MySQL InnoDB 在 RR 级别下通过 MVCC + Gap Lock（间隙锁）解决幻读问题：快照读（普通 SELECT）通过 MVCC 保证一致性读；当前读（SELECT FOR UPDATE / INSERT / UPDATE）通过 Next-Key Lock（Record Lock + Gap Lock）防止其他事务插入。')

# ==================== 三、PostgreSQL ====================
SEC2('三、PostgreSQL 数据库')
SEC3('3.1 索引与查询优化')

Q('Q：PostgreSQL 支持哪些索引类型？')
TBL(['索引类型','适用场景','说明'],[
    ['B-tree','等值、范围查询（默认）','最通用，支持 <, <=, =, >=, >, BETWEEN, IN, IS NULL'],
    ['Hash','等值查询','只支持 =，通常 B-tree 更好（PG 10 前 Hash 不持久化，崩溃后需重建）'],
    ['GiST','几何、全文搜索、范围类型','Generalized Search Tree，支持多种数据类型的复杂查询'],
    ['GIN','全文搜索、数组、JSONB','Generalized Inverted Index，倒排索引，适合多值属性'],
    ['BRIN','物理有序的大表（时间序列）','Block Range Index，极小的索引体积，适合自然有序的大表'],
    ['Bloom','多列等值查询','基于 Bloom Filter，空间效率高但有假阳性'],
])

A('追问 1：什么是最左前缀匹配？')
ANS('复合索引（多列索引）遵循最左前缀原则：索引 (a, b, c) 可以被以下查询使用：')
ABL('WHERE a = 1 —— 使用索引（匹配第一列）')
ABL('WHERE a = 1 AND b = 2 —— 使用索引（匹配前两列）')
ABL('WHERE a = 1 AND b = 2 AND c = 3 —— 使用索引（匹配全部三列）')
ABL('WHERE b = 2 —— 不使用索引（跳过了第一列 a）')
ABL('WHERE a = 1 AND c = 3 —— 部分使用索引（只用到 a，c 无法使用因为中间的 b 被跳过）')
ABL('WHERE a = 1 AND b > 2 AND c = 3 —— a 等值 + b 范围可以用索引，但 c 无法使用（范围查询后的列无法使用 B-tree 索引）')
ANS('设计原则：把等值查询的列放在前面，范围查询的列放在后面。把选择性高（区分度大）的列放在前面。')

A('追问 2：EXPLAIN ANALYZE 如何看？')
ANS('EXPLAIN (ANALYZE, BUFFERS, FORMAT TEXT) 输出两部分：')
ABL('预估计划（EXPLAIN 部分）：展示查询计划树，从内向外阅读。关注：Seq Scan（全表扫描，大表应避免）、Index Scan / Index Only Scan（索引扫描，理想状态）、Hash Join / Nested Loop / Merge Join（连接方式）、Sort（是否利用了索引排序）。')
ABL('实际执行（ANALYZE 部分）：每个节点的实际执行时间（actual time=首次返回时间..总时间）和实际行数（rows）。对比预估行数（rows=）和实际行数，如果偏差很大（如预估 1 行实际 10000 行），说明统计信息过旧，需要执行 ANALYZE table_name 更新统计信息。')
ANS('Buffers 部分：shared hit（缓存命中）vs shared read（磁盘读取），hit 越多越好。如果 read 很多，考虑增加 shared_buffers 配置。')
ANS('常见性能问题信号：Seq Scan on 大表、actual rows 远大于 estimated rows、Sort Method: external merge（内存不足需要磁盘排序）。')

A('追问 3：覆盖索引是什么？')
ANS('覆盖索引（Covering Index）是指索引中包含了查询所需的所有列，数据库可以直接从索引中返回结果，无需回表查询（Heap Fetch）。在 PostgreSQL 中通过 INCLUDE 子句实现：')
CODE('CREATE INDEX idx_cover ON orders (tenant_id, created_at) INCLUDE (amount, status);')
ANS("执行 SELECT amount, status FROM orders WHERE tenant_id = 1 AND created_at > '2024-01-01' 时，所有需要的数据都在索引中，EXPLAIN 显示 Index Only Scan，性能大幅提升。")
ANS('适用场景：高频查询返回的列较少且固定的场景。代价：索引体积增大，写入时维护索引的开销增加。')

SEC3('3.2 PostgreSQL 高级特性')

Q('Q：如何优化慢 SQL？（系统性回答）')
P('排查五步法：', bold=True, sz=10)
B('1. 定位慢 SQL：开启 log_min_duration_statement = 500（记录超过 500ms 的 SQL），或使用 pg_stat_statements 扩展统计 Top N 慢查询。')
B('2. 分析执行计划：EXPLAIN (ANALYZE, BUFFERS, FORMAT TEXT)。')
B('3. 检查索引使用：确认是否走了索引（Seq Scan 大表 = 问题）。')
B('4. 检查统计信息：ANALYZE table_name 更新统计信息（预估行数偏差大时）。')
B('5. 检查隐式类型转换：WHERE varchar_column = 123 会导致索引失效（类型不匹配时 PG 会在列上做类型转换函数）。')
P('常见优化手段：', bold=True, sz=10)
B('添加合适的索引（B-tree / GIN / BRIN 根据查询模式选择）。')
B('避免 SELECT *：只查需要的列，减少 IO 和网络传输。')
B("避免在 WHERE 中对列做函数运算：WHERE date(created_at) = '2024-01-01' 会失效索引，改写为 WHERE created_at >= '2024-01-01' AND created_at < '2024-01-02'。")
B('大分页优化：OFFSET 100000 LIMIT 10 会扫描并丢弃前 100000 行。改用 Keyset Pagination：WHERE id > last_seen_id ORDER BY id LIMIT 10。')
B('子查询改 JOIN：相关子查询可能导致逐行执行，改写为 JOIN 后优化器可以有更好的选择。')
B('CTE 优化：PG 12+ 的 CTE 默认是内联的（优化器可以跨 CTE 优化），使用 MATERIALIZED 关键字强制物化（防止优化器选择错误的计划时）。')

Q('Q：PostgreSQL MVCC 原理？')
ANS('PostgreSQL 的 MVCC（多版本并发控制）通过在每行数据中维护多个版本来实现无读锁的并发控制：')
ABL('每行数据有 xmin（创建该行的事务 ID）和 xmax（删除/更新该行的事务 ID，0 表示未删除）。')
ABL('UPDATE 操作不是原地修改，而是：标记旧行的 xmax 为当前事务 ID + INSERT 一行新版本（Heap Tuple）。这就是为什么 PostgreSQL 的 UPDATE 比 MySQL 更慢（需要写更多数据），但也避免了 Undo Log 的复杂性。')
ABL('DELETE 操作只设置 xmax，物理删除由 VACUUM 负责。')
ABL('SELECT 操作通过事务快照（Snapshot）判断每行是否对当前事务可见：只读取 xmin 已提交且 xmax 未提交（或 xmax 为 0）的行。不需要加任何读锁，读写完全不阻塞。')
ANS('需要 VACUUM 的原因：UPDATE/DELETE 产生的旧版本（dead tuples）不会自动清理，需要 VACUUM 进程定期扫描并回收空间。Autovacuum 是 PG 内置的自动 VACUUM 进程，通过 autovacuum_vacuum_threshold、autovacuum_vacuum_scale_factor 等参数控制触发频率。')
ANS("表膨胀（Table Bloat）：如果长事务（持有旧快照）阻止 VACUUM 清理 dead tuples，表的物理大小会持续增长，查询性能下降。检测方法：pgstattuple 扩展的 pgstattuple('table_name') 查看 dead_tuple_percent。处理：(1) 终止长事务；(2) VACUUM FULL（需要排他锁，会锁表）；(3) pg_repack（在线重建，不锁表，推荐）。")

# ==================== 四、Redis ====================
SEC2('四、Redis')

Q('Q：Redis 常用数据结构和使用场景？')
TBL(['类型','底层编码','典型使用场景'],[
    ['String','SDS / int / embstr','缓存对象(JSON)、计数器(INCR)、分布式锁(SET NX)、Session'],
    ['List','quicklist(ziplist+linkedlist)','消息队列(LPUSH+BRPOP)、最新列表、时间线'],
    ['Hash','ziplist / hashtable','对象缓存(用户信息)、购物车、配置存储'],
    ['Set','intset / hashtable','标签系统、共同好友(SINTER)、抽奖(SRANDMEMBER/SPOP)'],
    ['ZSet','ziplist / skiplist+hashtable','排行榜(ZREVRANGE)、延迟队列(ZADD+ZPOPMIN)、滑动窗口限流'],
    ['Bitmap','String','签到统计(GETBIT/SETBIT)、日活统计(BITOP OR)、在线状态'],
    ['HyperLogLog','String','UV 统计（误差 0.81%，固定 12KB 内存，不管多少数据）'],
    ['Stream','rax(listpack)','消息队列（类 Kafka，支持消费者组、ACK、持久化）'],
])

SEC3('4.2 缓存三大问题')

Q('Q：缓存穿透、击穿、雪崩的区别和解决方案？')
TBL(['问题','现象','根因','解决方案'],[
    ['穿透','大量请求查询不存在的 key，每次都打 DB','恶意攻击或代码 Bug（查询不存在的数据）','1. 布隆过滤器拦截不存在的 key\n2. 缓存空值（短 TTL，如 60s）\n3. 接口层参数校验'],
    ['击穿','热点 key 过期瞬间，大量并发请求同时打 DB','热点 key 过期 + 高并发','1. 互斥锁（SETNX），只允许一个线程回源 DB\n2. 逻辑过期：value 中存储过期时间，后台异步刷新\n3. 热点 key 永不过期 + 后台定时刷新'],
    ['雪崩','大面积 key 同时过期 / Redis 宕机，请求全部打 DB','大批 key 设置相同 TTL / Redis 故障','1. TTL 加随机值（base_ttl + random(0,300)）\n2. Redis 集群高可用（Sentinel / Cluster）\n3. 本地缓存（Caffeine/Guava）兜底\n4. 限流降级（Sentinel/Hystrix）'],
])

A('追问：缓存穿透的布隆过滤器原理？')
ANS('布隆过滤器（Bloom Filter）是一种空间效率极高的概率型数据结构，用于判断一个元素是否「可能在集合中」或「一定不在集合中」。')
ANS('原理：底层是一个 bit 数组（如 1024 位）和 k 个独立的哈希函数。添加元素时，用 k 个哈希函数分别计算得到 k 个位置，将这些位置的 bit 设为 1。查询元素时，检查这 k 个位置是否全为 1：如果有一位为 0，则元素一定不存在（100% 准确）；如果全为 1，则元素可能存在（有假阳性，因为不同的元素可能碰巧设置相同的位）。')
ANS('在缓存穿透场景中的使用：在 Redis 查询之前，先用布隆过滤器判断 key 是否存在。如果布隆过滤器说不存在，直接返回，不查 DB。Redis 中可以使用 RedisBloom 模块的 BF.ADD / BF.EXISTS 命令。')
ANS('注意事项：(1) 布隆过滤器不支持删除（删除会影响其他元素），如需删除使用 Counting Bloom Filter 或 Cuckoo Filter；(2) 需要预估元素数量和可接受的误判率来计算 bit 数组大小和哈希函数个数。')

SEC3('4.3 分布式锁')

Q('Q：Redis 分布式锁如何实现？有什么问题？')
CODE('''-- 加锁（原子操作：SET key value NX PX 30000）
SET lock_key <unique_id> NX PX 30000

-- 释放锁（Lua 脚本保证原子性：先比较 value 再删除）
if redis.call('get', KEYS[1]) == ARGV[1] then
    return redis.call('del', KEYS[1])
else
    return 0
end''')
ANS('为什么 value 用 unique_id？防止误删其他线程的锁：线程 A 获取锁后执行超时（超过 30s），锁自动过期释放；线程 B 获取到锁；线程 A 执行完后如果不检查 value 就删除，会把线程 B 的锁删掉。unique_id 可以用 UUID 生成。')

A('追问 1：锁超时但业务未完成怎么办？')
ANS('问题：如果业务执行时间超过锁的 TTL（如 30s），锁会自动释放，其他线程可能获取到锁，导致并发问题。')
ANS('解决方案——看门狗（Watchdog）机制：Redisson 客户端内置了看门狗机制。获取锁成功后，启动一个后台定时任务，每隔 lockWatchdogTimeout/3（默认 10 秒）检查锁是否还被持有，如果还在则自动续期（将 TTL 重置为 30s）。只有在不指定 leaseTime 时才启用看门狗。')
CODE('''// Redisson 自动续期\nRLock lock = redisson.getLock(\"myLock\");\nlock.lock();  // 不指定 leaseTime，看门狗自动续期\ntry {\n    // 业务逻辑，执行多久都不会过期\n} finally {\n    lock.unlock();  // 释放锁，同时停止看门狗\n}''')

A('追问 2：主从切换丢锁问题？')
ANS('场景：(1) 线程 A 在 Master 上获取锁成功；(2) Master 宕机，锁数据未同步到 Slave；(3) Slave 升级为新 Master；(4) 线程 B 在新 Master 上获取锁成功 —— 两个线程同时持有了锁。')
ANS('解决方案——RedLock 算法（Redis 作者 Antirez 提出）：向 N 个（推荐 5 个）独立的 Redis 实例分别加锁，如果在 N/2+1 个以上实例加锁成功且总耗时小于锁的 TTL，则认为加锁成功。')
ANS('争议：Martin Kleppmann 在《How to do distributed locking》中指出 RedLock 的问题：(1) 依赖时钟同步（GC 暂停、时钟跳跃可能导致锁提前过期）；(2) 没有 fencing token 机制防止过期锁的操作重叠。建议使用共识算法（ZooKeeper/etcd）代替 RedLock。')
ANS('实际建议：如果对一致性要求不高（如防重复提交），单实例 Redis 锁够用；如果要求强一致性，使用 ZooKeeper（临时顺序节点 + watch）或 etcd（lease + revision）。')

SEC3('4.4 Redis 持久化与集群')

Q('Q：RDB vs AOF？')
TBL(['维度','RDB','AOF'],[
    ['方式','定时 fork 子进程生成全量快照（二进制）','追加写命令到日志文件'],
    ['触发','save 900 1 / save 300 10 / BGSAVE','always / everysec / no'],
    ['恢复速度','快（直接加载二进制）','慢（需要重放所有命令）'],
    ['数据安全','可能丢失最后一次快照后的数据','everysec 最多丢 1 秒数据'],
    ['文件大小','紧凑（压缩二进制）','较大（文本命令，需要 rewrite 压缩）'],
    ['fork 开销','fork 时需要 Copy-On-Write，大内存时耗时','AOF rewrite 也需要 fork'],
])
ANS('Redis 4.0+ 混合持久化（aof-use-rdb-preamble yes）：AOF rewrite 时先以 RDB 格式写入全量数据，增量部分以 AOF 格式追加。这样恢复时先加载 RDB 部分（快），再重放 AOF 部分（少量命令），兼顾了速度和安全性。')

Q('Q：Redis Cluster 原理？')
ANS('Redis Cluster 是 Redis 的分布式集群方案，核心设计：')
ABL('数据分片：使用 16384 个哈希槽（slot），每个 key 通过 CRC16(key) % 16384 计算所属 slot，不同 slot 分配到不同节点。')
ABL('节点通信：使用 Gossip 协议（PING/PONG 消息）传播集群状态。每个节点知道所有 slot 到节点的映射关系。')
ABL('请求路由：客户端发送命令到任意节点，如果 key 不在该节点，返回 MOVED 重定向（永久重定向）或 ASK 重定向（临时重定向，slot 迁移中）。智能客户端（如 JedisCluster/Lettuce）会缓存 slot-node 映射，直接路由到正确节点。')
ABL('故障检测：每个节点定期向其他节点发送 PING，超时未收到 PONG 则标记为 PFAIL（主观下线）。当半数以上 master 节点都标记某节点 PFAIL 时，升级为 FAIL（客观下线），触发该节点的 slave 发起选举。')
ABL('故障转移：客观下线的 master 的 slave 发起 Raft 选举，获得多数投票后升级为新 master，接管原 master 的 slot。')
ANS('Cluster 限制：(1) 不支持跨 slot 的多 key 操作（如 MGET 不同 slot 的 key），需要使用 Hash Tag（{user}:1001 和 {user}:1002 强制同 slot）；(2) 只能使用 db0。')

# ==================== 五、Kafka ====================
SEC2('五、Kafka')

Q('Q：如何保证消息不丢失？')
ANS('消息丢失可能发生在三个环节，需要分别保证：')
P('Producer 端：', bold=True, sz=10)
ABL('acks=all（或 acks=-1）：要求所有 ISR（In-Sync Replicas）副本都确认写入后才返回成功。acks=0 表示不等确认（最快但可能丢消息），acks=1 表示 Leader 确认（Leader 宕机可能丢）。')
ABL('retries > 0 + retry.backoff.ms：发送失败时自动重试。')
ABL('min.insync.replicas=2：至少需要 2 个 ISR 副本，如果 ISR 数量不足，Producer 会收到 NotEnoughReplicasException，拒绝发送。')
P('Broker 端：', bold=True, sz=10)
ABL('replication.factor >= 3：每个分区至少 3 个副本。')
ABL('unclean.leader.election.enable=false：禁止非 ISR 副本成为 Leader（防止数据丢失）。如果所有 ISR 副本都宕机，宁可服务不可用也不要选一个数据不完整的 Follower 当 Leader。')
P('Consumer 端：', bold=True, sz=10)
ABL('enable.auto.commit=false：关闭自动提交 offset。如果自动提交，Consumer 拉取消息后还没处理就提交了 offset，此时 Consumer 崩溃，消息就丢了。')
ABL('手动提交：处理完消息后手动调用 commitSync() 或 commitAsync()。')

Q('Q：如何保证消息有序？')
ANS('Kafka 只保证单个 Partition 内的消息有序（按写入顺序追加到 Log），不保证跨 Partition 的全局有序。')
ANS('方案一（全局有序）：Topic 只设 1 个分区。所有消息写入同一个 Partition，天然有序。缺点：无法水平扩展（只能单个 Consumer 消费），吞吐受限。')
ANS('方案二（业务有序，推荐）：相同业务 key 的消息发到同一个 Partition。Producer 发送时指定 key（如订单 ID），Kafka 对 key 做 hash 决定 Partition。这样同一个订单的所有状态变更消息都在同一个 Partition 中，保证顺序。不同订单之间并行消费，兼顾吞吐。')
CODE('// 指定 key 保证业务有序\nproducer.send(new ProducerRecord<>(\"order-events\", orderId, eventData));')

A('追问：消费者 Rebalance 机制？')
ANS('Rebalance 是 Consumer Group 中分区重新分配的过程，触发条件：(1) 新 Consumer 加入组；(2) Consumer 离开组（主动退出或心跳超时被踢出）；(3) Consumer 订阅的 Topic 分区数变化。')
ANS('Rebalance 的问题：(1) Rebalance 期间所有 Consumer 暂停消费（Stop-The-World），导致消费延迟；(2) Rebalance 耗时可能很长（默认 session.timeout.ms=45s），特别是消费者多的时候。')
ANS('优化方案：')
ABL('缩短检测时间：减小 session.timeout.ms 和 heartbeat.interval.ms（但太小会导致网络抖动误判）。')
ABL('静态成员（Static Membership，Kafka 2.3+）：Consumer 配置 group.instance.id，短暂重启时不会触发 Rebalance（只要 instance.id 不变且在 session.timeout.ms 内回来）。')
ABL('Cooperative Rebalance（增量式，Kafka 2.4+）：将 eager rebalance（全部停止-重分配-恢复）改为 cooperative（只重分配变化的分区，未变化的分区继续消费）。partition.assignment.strategy=org.apache.kafka.clients.consumer.CooperativeStickyAssignor。')

# ==================== 六、MyBatis ====================
SEC2('六、MyBatis')

Q('Q：#{} 和 ${} 的区别？')
TBL(['维度','#{}','${}'],[
    ['处理方式','预编译参数替换（PreparedStatement ?）','字符串直接拼接'],
    ['SQL 注入','安全（参数被转义）','不安全（直接拼入 SQL）'],
    ['使用场景','参数值（WHERE 条件值、INSERT 值）','表名、列名、ORDER BY 子句等不能用 ? 的地方'],
    ['示例','WHERE id = #{id} -> WHERE id = ?','ORDER BY ${column} -> ORDER BY create_time'],
])

A('追问：什么时候必须用 ${}？')
ANS('当 SQL 的结构部分（不是值部分）需要动态化时，必须用 ${}：')
ABL('动态表名：SELECT * FROM ${tableName} WHERE id = #{id}')
ABL('动态列名：SELECT ${columns} FROM users')
ABL('动态 ORDER BY：ORDER BY ${sortField} ${sortOrder}')
ANS('使用 ${} 时必须在代码层做白名单校验，不能直接使用用户输入。')

Q('Q：MyBatis 一级缓存和二级缓存？')
TBL(['维度','一级缓存','二级缓存'],[
    ['范围','SqlSession 级别','Mapper（namespace）级别'],
    ['默认','开启','关闭（需 <cache/> 配置）'],
    ['失效条件','写操作（INSERT/UPDATE/DELETE）、clearCache()、SqlSession 关闭','同 namespace 的写操作'],
    ['Spring 中','几乎无效：Spring 每次请求创建新 SqlSession，SqlSession 之间缓存不共享','谨慎使用：多表关联查询时缓存可能不一致'],
])
ANS('Spring + MyBatis 中一级缓存几乎无效的原因：Spring 通过 SqlSessionTemplate 管理 SqlSession，每次 Mapper 方法调用结束后就关闭 SqlSession（归还到 SqlSessionHolder），下次调用创建新 SqlSession，一级缓存自然失效。')
ANS('二级缓存的坑：namespace 级别意味着如果 Mapper A 查询了表 user 和表 order（JOIN 查询），Mapper B 单独更新了 order 表，Mapper A 的缓存不会失效。这导致脏数据。建议：在高一致性场景下关闭二级缓存，使用 Redis 等外部缓存自行管理。')

Q('Q：MyBatis 的 N+1 问题？')
ANS('N+1 问题发生在嵌套查询（嵌套结果映射 select 属性）中：')
CODE('''<!-- 1 次查询得到 N 条订单，每条订单触发 1 次查询查 items => N+1 次 -->\n<resultMap id="orderMap" type="Order">\n    <id property="id" column="id"/>\n    <collection property="items" ofType="Item"\n                select="getItemsByOrderId" column="id"/>\n</resultMap>''')
ANS('解决方案：')
ABL('方案一（推荐）：嵌套结果映射（JOIN），一次 SQL 查完：')
CODE('''<resultMap id="orderMap" type="Order">\n    <id property="id" column="id"/>\n    <collection property="items" ofType="Item">\n        <id property="id" column="item_id"/>\n        <result property="name" column="item_name"/>\n    </collection>\n</resultMap>\n<select id="getOrderWithItems">\n    SELECT o.*, i.id item_id, i.name item_name\n    FROM orders o LEFT JOIN items i ON o.id = i.order_id\n</select>''')
ABL('方案二：批量查询，两步走：先查出所有 order id，再用 IN 一次查所有 items，代码中手动组装。')
ABL('方案三：MyBatis 延迟加载（lazyLoadingEnabled=true），只在访问 items 属性时才触发查询。但要注意 SqlSession 生命周期，SqlSession 关闭后无法延迟加载。')

print('Part 2 done')
doc.save(OUTPUT)
print(f'Saved {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')
