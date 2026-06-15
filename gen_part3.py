"""Part 3: Project deep dives - all 4 projects with full answers"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = r'C:\Users\47583\projects\hisi_dev_tool v5.0\沈江春_面试准备手册_完整版.docx'
doc = Document(OUTPUT)

def P(text, bold=False, sz=None, clr=None, align=None, sa=None, sn=None):
    p = doc.add_paragraph(style=sn) if sn else doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if sz: r.font.size = Pt(sz)
    if clr: r.font.color.rgb = RGBColor.from_string(clr)
    if align: p.alignment = align
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    return p
def B(text, lv=0):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    r = p.add_run(text); r.font.size = Pt(10)
    if lv: p.paragraph_format.left_indent = Cm(1.5*lv)
def TBL(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        c.paragraphs[0].add_run(h).bold = True; c.paragraphs[0].runs[0].font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            c.paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()
def CODE(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
def A(title): P(title, bold=True, sz=11, clr='2E75B6', sa=2)
def ANS(text): P(text, sz=10, sa=6)
def ABL(text): B(text)
def Q(text): P(text, bold=True, sz=11, sa=4)
def SEC1(t): doc.add_heading(t, level=1)
def SEC2(t): doc.add_heading(t, level=2)
def SEC3(t): doc.add_heading(t, level=3)

# ============================================================
# PART 2: 项目深挖
# ============================================================
SEC1('第二部分：项目深挖 Q&A（含完整答案）')

# ====== 项目一：纳税申报 SaaS ======
SEC2('项目一：纳税申报 SaaS 平台')

Q('Q1：请介绍纳税申报 SaaS 平台的架构？')
B('单体 Spring Boot 应用，面向 300+ 子公司的多租户架构。')
B('多租户隔离方式：共享数据库 + 租户 ID 字段隔离（tenant_id），通过 MyBatis 拦截器自动注入 tenant_id 条件。')
B('核心模块：纳税申报（填写/提交/审核）、税款缴纳（对接银行支付）、纳税报表（生成/导出）、税收计算引擎（多国税率规则）。')
B('支持中国、马来西亚、新加坡、印尼、日本等多国税务规则。')
B('技术栈：Spring Boot 3 / JDK 21 / PostgreSQL / Redis / Kafka / MyBatis。')
B('外部对接：税务局接口、银行支付接口、ERP 系统数据同步。')

Q('Q2：多租户架构有哪些实现方式？你项目用的哪种？')
TBL(['方式','隔离级别','成本','适用场景'],[
    ['独立数据库','最高（物理隔离）','高（每个租户一个数据库实例）','安全性要求极高（如金融、医疗）'],
    ['共享数据库+独立 Schema','中（逻辑隔离）','中（每个租户一个 Schema）','租户数量适中（几十到几百）'],
    ['共享数据库+tenant_id 列','最低（应用层隔离）','低（所有租户共享表）','租户数量多（300+），成本敏感'],
])

A('追问 1：如何保证不查询到其他租户的数据？')
ANS('我们在应用层实现了自动化的租户隔离，核心机制：')
ABL('全局 ThreadLocal 存储当前请求的 tenant_id：在 Filter/Interceptor 层从请求头或 JWT Token 中解析 tenant_id，存入 TenantContext（ThreadLocal 封装类）。')
ABL('MyBatis 拦截器自动注入：实现 Interceptor 接口，在 sqlSessionFactory 中注册。拦截所有 SELECT/UPDATE/DELETE 的 SQL，自动追加 WHERE tenant_id = #{tenantId} 条件。这样开发人员写 SQL 时不需要手动加 tenant_id 条件，减少遗漏风险。')
ABL('INSERT 时自动填充 tenant_id：通过 MyBatis MetaObjectHandler 在 INSERT 时自动设置 tenant_id 字段值。')
ANS('防御措施：(1) 每个请求入口强制校验 tenant_id 有效性；(2) 定期安全审计，扫描是否有 SQL 绕过拦截器直接执行的情况；(3) 数据库层面的 Row Level Security（RLS）作为第二道防线（PostgreSQL 支持）。')

A('追问 2：租户间性能隔离如何做？')
ANS('在共享数据库模式下，一个租户的高负载可能影响其他租户（noisy neighbor 问题）。我们的做法：')
ABL('API 限流：基于 Redis + 滑动窗口实现每租户的 API 调用频率限制，超限返回 429。')
ABL('数据库连接池限制：HikariCP 按租户维度设置最大连接数（通过动态数据源路由），防止单个租户耗尽连接池。')
ABL('查询超时：PostgreSQL 的 statement_timeout 参数设置单条 SQL 的最大执行时间（如 30s），防止慢查询拖垮整个数据库。')
ABL('资源配额管理：大报表查询走异步任务队列，限制并发生成数，避免大量复杂查询同时执行。')

Q('Q3：多国税务计算引擎如何设计？')
ANS('采用策略模式 + 模板方法模式设计，核心思想是将「变化的」（各国税务规则）和「不变的」（计算流程）分离：')
B('1. 定义统一接口 TaxCalculator：calculateTax(ReportData) -> TaxResult。')
B('2. 定义模板抽象类 AbstractTaxCalculator，封装通用流程：validateData -> getTaxRates -> calculateByType -> generateReport。')
B('3. 每个国家实现自己的子类：ChinaTaxCalculator、MalaysiaTaxCalculator、SingaporeTaxCalculator 等。')
B('4. 通过工厂模式根据 tenant 的国家配置自动选择对应的 Calculator 实现。')
B('5. 税率和计算规则通过数据库配置管理，支持动态更新而不需要改代码发版。')

A('追问：如何保证计算精度？')
ANS('财务计算必须精确到分，绝对不能使用 float/double（IEEE 754 浮点数有精度问题，如 0.1 + 0.2 != 0.3）。具体做法：')
ABL('Java 代码中全链路使用 BigDecimal，构造时用字符串：new BigDecimal(\"0.1\") 而不是 new BigDecimal(0.1)。')
ABL('BigDecimal 的除法指定舍入模式：divide(divisor, 2, RoundingMode.HALF_UP)（保留 2 位小数，四舍五入）。')
ABL('PostgreSQL 中使用 NUMERIC(precision, scale) 类型存储（如 NUMERIC(18, 2) 表示最多 18 位数字，2 位小数），而不是 FLOAT/DOUBLE。')
ABL('前后端传输金额时使用分为单位的 long 类型（避免 JSON 序列化的精度问题），展示时由前端除以 100 并格式化。')
ABL('对账机制：每日自动对账，比对申报系统计算的税额和税务局接口返回的税额，差异超过阈值自动告警。')

Q('Q4：你排查的内存泄漏具体是什么情况？')
ANS('用 STAR 法则完整描述：')
ABL('Situation：纳税申报系统在月底报税高峰期（每月 25-30 号），运行 2-3 天后响应时间从正常的 200ms 逐步上升到 2-5s，同时 GC 日志显示 Full GC 频率从每天 1-2 次增加到每小时多次。')
ABL('Task：需要在不影响线上业务的情况下（报税高峰期不能停服），定位根因并制定修复方案。')
ABL('Action：第一步，通过 jstat -gcutil <pid> 1000 观察到 Full GC 后老年代占用从 40% 逐步上升到 85%，确认存在内存泄漏。第二步，在凌晨低峰期执行 jmap -dump 导出堆转储文件（约 3GB）。第三步，使用 MAT 分析 Dominator Tree，发现一个 ConcurrentHashMap（报表缓存）持有约 1.2GB 内存，key 是 tenantId + reportType + month 的组合。第四步，定位代码发现报表计算模块为避免重复计算，将每次计算的中间结果（包含大量 BigDecimal 计算中间值和数据列表）缓存到一个 static ConcurrentHashMap 中，但没有设置容量上限和过期策略。随着租户数和报表类型增多，缓存无限增长。')
ABL('Result：引入 Caffeine 本地缓存替换手动 HashMap，配置 maximumSize=500（最多缓存 500 个报表结果）和 expireAfterAccess=30min（30 分钟未访问自动过期）。上线后观察一周，老年代使用率稳定在 30-50%，Full GC 恢复到每天 1-2 次，响应时间恢复到 200ms 以内。')

Q('Q5：SQL 性能问题如何排查和优化？')
ANS('准备具体案例：')
ABL('发现：报表查询接口在数据量大的月份（如年度汇总报表）响应时间超过 10 秒，慢查询日志中出现大量相关 SQL。')
ABL('排查：EXPLAIN (ANALYZE, BUFFERS) 分析发现该 SQL 执行了 Seq Scan on tax_report（全表扫描），扫描行数约 500 万行，但实际返回只有 12 行。原因：WHERE 条件中使用了函数：WHERE EXTRACT(YEAR FROM report_date) = 2024 AND tenant_id = #{tenantId}，PostgreSQL 无法对 EXTRACT(YEAR FROM report_date) 使用 report_date 列上的 B-tree 索引。')
ABL("优化方案：(1) 改写查询条件为范围查询：WHERE report_date >= '2024-01-01' AND report_date < '2025-01-01' AND tenant_id = #{tenantId}，这样可以使用 (tenant_id, report_date) 复合索引。(2) 同时创建覆盖索引：CREATE INDEX idx_tax_report_cover ON tax_report(tenant_id, report_date) INCLUDE(amount, status)。")
ABL('结果：查询时间从 10s+ 降到 150ms，Index Only Scan 替代了 Seq Scan，buffer 命中率从 2% 提升到 99%。')

Q('Q6：如何保证财务数据的准确性？')
ANS('财务系统对数据准确性要求极高，我们从多个层面保证：')
ABL('计算精度：全链路 BigDecimal + PostgreSQL NUMERIC 类型，禁止 float/double。除法操作统一指定 RoundingMode.HALF_UP。')
ABL('幂等设计：相同申报单重复提交不会重复扣款。通过申报单唯一编号（幂等键）+ 数据库唯一约束 + Redis 分布式锁三重保障。')
ABL('对账机制：每日自动对账 Job，比对三个维度：(1) 申报系统税额 vs 税务局接口返回税额；(2) 系统记录的缴税金额 vs 银行实际扣款金额；(3) 日终汇总 vs 分户明细。差异超过阈值自动告警到值班群。')
ABL('审计日志：所有财务数据变更操作（创建/修改/提交/审核/驳回）记录完整的操作日志，包括操作人、操作时间、变更前后值（JSON diff）。使用数据库触发器 + 独立审计表，确保日志不可篡改。')
ABL('四眼原则：关键财务操作（如大额税款缴纳）需要两级审批，系统强制校验。')

Q('Q7：Kafka / Redis 在纳税系统中的使用场景？')
ANS('Kafka 使用场景：')
ABL('异步通知：申报完成后通过 Kafka 发送通知事件，下游的邮件服务、短信服务消费事件后发送通知。解耦了申报核心流程和通知流程。')
ABL('削峰填谷：月底申报高峰期，前端提交的申报请求先进入 Kafka Topic，后端消费者按自己的处理能力消费。保护数据库不被瞬时高流量击垮。')
ABL('事件驱动：申报状态变更（草稿->提交->审核->通过->缴税）通过 Kafka 事件通知 ERP、财务报表等下游系统，实现系统间解耦。')
ANS('Redis 使用场景：')
ABL('Session 存储：多实例部署时共享用户登录态，支持会话级别的租户信息缓存。')
ABL('热点数据缓存：税率配置（变动频率低但查询频繁）、公司基础信息、申报模板等。')
ABL('分布式锁：防重复提交申报（同一租户同一报表类型在一定时间内只能有一个提交请求在处理）。')
ABL('限流计数：基于 Redis 的滑动窗口限流，保护 API 不被恶意调用。')

# ====== 项目二：研发管理工具链 ======
SEC2('项目二：研发管理工具链')

Q('Q1：开源三方件安全整改的具体流程？')
ANS('完整的安全整改流程分为五个阶段：')
B('1. 扫描阶段：使用 OWASP Dependency-Check 和 Snyk 对项目的 Maven 依赖树进行全量扫描。扫描工具会对比 CVE 数据库（NVD），识别已知漏洞。')
B('2. 梳理阶段：整理扫描结果，共发现 246 个组件中有 41 个高危 CVE（CVSS >= 7.0）。按组件分类，标注影响范围（直接依赖 vs 传递依赖）、漏洞类型（RCE/反序列化/SQL注入等）、修复版本。')
B('3. 评估阶段：对每个 CVE 评估实际风险：(1) 漏洞是否在我们的使用场景中可被触发（如 Spring 的某个漏洞需要特定配置才可利用）；(2) 是否有替代方案；(3) 升级版本的兼容性影响。输出风险评估报告。')
B('4. 修复阶段：三种修复策略——(a) 升级版本：如 log4j 2.x 升级到 2.17+（修复 Log4Shell RCE 漏洞）；(b) 替换组件：如将不维护的 commons-collections 3.x 替换为 4.x；(c) 排除传递依赖：在 pom.xml 中用 <exclusion> 排除不需要的传递依赖。同时推动核心框架升级：JDK 8 -> 21、Spring 2.x -> 3.4。')
B('5. 验证阶段：修复后重新扫描确认漏洞消除。全量回归测试确保功能不受影响。建立持续扫描机制（CI/CD pipeline 中集成 Dependency-Check，每次构建自动扫描）。')

Q('Q2：JDK 8 升 21 的主要挑战？')
TBL(['挑战','具体问题','解决方案'],[
    ['javax -> jakarta','Spring Boot 3 强制要求 Jakarta EE 9+，所有 javax.servlet/javax.persistence 等包名变更','全局替换 + OpenRewrite 自动化工具 + 三方库升级到 Jakarta 兼容版本'],
    ['反射访问限制','JDK 16+ 默认禁止非法反射（--illegal-access=deny）','对必要的反射添加 --add-opens 参数，优先重构代码消除反射依赖'],
    ['Security Manager','JDK 17 废弃，JDK 18 移除相关 API','移除所有 SecurityManager 相关代码，改用其他安全方案'],
    ['GC 参数变化','CMS 被移除（JDK 14），相关 JVM 参数无效','迁移到 G1（默认）或 ZGC，调整 GC 相关参数'],
    ['三方库兼容性','部分老版本库不支持 JDK 21（如 ASM 版本不够新）','升级三方库到支持 JDK 21 的版本，必要时排除传递依赖后手动指定'],
    ['Nashorn 移除','JDK 15 移除了 JS 引擎','迁移脚本到 GraalVM JS 或移除 JS 执行需求'],
])

A('追问：升级过程中具体怎么做的？分了几个阶段？')
ANS('我们分为四个阶段推进：')
B('阶段一（1 周）：评估。梳理所有依赖的 JDK 兼容性，输出影响范围文档。建立升级工作组（3 人），明确分工。')
B('阶段二（2 周）：基础框架升级。先升级 Spring Boot 2.x -> 3.4 + JDK 8 -> 21，处理 javax -> jakarta 迁移。在独立分支开发，不影响主线。')
B('阶段三（2 周）：三方库兼容。逐一升级 246 个组件中受影响的版本，处理传递依赖冲突。每升级一个组件就跑全量单元测试。')
B('阶段四（2 周）：集成测试 + 灰度。全量功能测试、性能测试（对比升级前后的吞吐和延迟）、安全扫描（确认 CVE 消除）。先在测试环境灰度，再逐步推到生产。')
ANS('踩过的坑：(1) 某些内部工具类使用了 sun.misc.Unsafe，JDK 21 中需要 --add-opens 参数；(2) Jackson 的某些序列化行为在 JDK 21 中有细微变化，导致已有数据的反序列化失败，需要配置兼容模式。')

Q('Q3：Spring 2.x 升 3.4 的主要挑战？')
B('1. Jakarta EE 命名空间迁移：javax.servlet.* -> jakarta.servlet.*，javax.persistence.* -> jakarta.persistence.*。这是最大的工作量，涉及所有 Controller、Service、Entity 类的 import 语句。')
B('2. Spring Security 配置方式变化：WebSecurityConfigurerAdapter 被移除，改用 @Bean SecurityFilterChain 方式。需要重写所有安全配置类。')
B('3. Spring MVC 路径匹配策略：AntPathMatcher -> PathPatternParser，某些复杂的 URL 匹配规则可能有差异，需要逐一验证。')
B('4. Micrometer + Observation API：Boot 3 内置了可观测性支持，原有的自定义 Metrics 代码需要适配新的 API。')
B('5. Hibernate 5 -> 6：JPA 实现从 Hibernate 5 升级到 6，HQL 语法和类型映射有细微变化。')

Q('Q4：如何推动跨团队的技术升级？')
ANS('技术升级最大的挑战不是技术本身，而是组织协调和风险管控：')
ABL('输出影响评估文档：列出所有受影响的组件、预估工作量、潜在风险和回滚方案。用数据说话（41 个高危 CVE 的 CVSS 评分和影响范围），让管理层理解升级的必要性。')
ABL('建立升级工作组：3 人核心团队 + 各模块 Owner 作为联络人。每周同步进展，及时暴露阻塞点。')
ABL('分阶段推进：先非核心模块验证可行性，积累经验和踩坑记录，再推广到核心模块。降低一次性风险。')
ABL('建立兼容性测试矩阵：每个三方库升级后跑全量自动化测试，确保不引入回归问题。')
ABL('输出升级指南：将过程中积累的经验和解决方案整理为文档，帮助其他团队复用。升级指南成为部门技术资产。')

# ====== 项目三：AI 研发提效 ======
SEC2('项目三：AI 研发提效专项')

Q('Q1：你开发了哪些 AI 辅助工具？具体怎么用的？')
ANS('三款核心工具：')
B('1. 代码提交追踪工具：在 Git Hook（post-commit）中触发，自动将 commit message 中的需求编号关联到项目管理系统中的具体需求。效果：需求-代码-测试的追溯链自动化，不需要手动维护 traceability matrix。')
B('2. 自动化测试录制回放：基于浏览器 Playwright 录制用户的 UI 操作流程，自动生成可回放的测试脚本。测试人员只需操作一遍系统，后续回归测试自动回放。效果：UI 回归测试从手动 2 天缩短到自动 30 分钟。')
B('3. 代码语义查询：基于 HiSi DevTool 的知识图谱和混合检索引擎，支持自然语言查询代码。例如查询「处理支付回调的方法」能精准定位到相关代码。效果：新人 onboarding 时间从 2 周缩短到 3 天。')

Q('Q2：Skill 套件是什么？如何覆盖 31 人？')
ANS('Skill 套件是基于 Claude Code 的 Skill 系统开发的团队级 AI 辅助工具集：')
ABL('技术原理：每个 Skill 是一个 .md 文件，定义了特定场景的操作流程和上下文信息。Claude Code 加载 Skill 后能按照预定义的流程执行任务。')
ABL('包含的 Skill：代码审查（check）、文档生成（init-docs）、测试用例生成（tdd）、提交规范检查（git-workflow）、需求分析（planning）等。')
ABL('分发方式：将 Skill 文件放在团队共享的 .claude/skills/ 目录下，通过 Git 仓库管理。每个开发者 clone 项目后自动获得所有 Skill。在 CLAUDE.md 中配置团队级规则和 Skill 引用。')
ABL('覆盖 31 人：团队 31 名开发者全部安装了 Claude Code，通过统一的配置文件（.claude/settings.json）确保所有人使用相同的 Skill 集合。使用率跟踪通过 Git 提交记录和 Skill 调用日志统计。')

Q('Q3：116 人技术沙龙讲了什么？效果如何？')
ANS('技术沙龙内容：')
ABL('主题：「AI 辅助开发的最佳实践——从工具到工作流」。')
ABL('内容模块：(1) AI 编程工具概览（Claude Code / Copilot / Cursor 对比）；(2) 实操演示：如何用 Skill 套件完成日常开发任务；(3) 效率数据：对比使用前后的开发效率（代码审查时间减少 40%，文档编写时间减少 60%）；(4) 踩坑分享：AI 工具的局限性和正确使用姿势。')
ABL('效果：(1) 两个产品团队（约 40 人）在沙龙后全面接入 AI 工具；(2) 部门月度例会中作为固定议题，每期分享一个 AI 使用技巧；(3) 累计 10 场分享后，团队 AI 工具使用率从 20% 提升到 85%。')

# ====== 项目四：HiSi DevTool（重点） ======
SEC2('项目四：HiSi DevTool — 代码知识图谱驱动的智能开发平台（重点准备）')
P('这是简历上最有技术深度的项目，面试官大概率会深挖。以下是每个模块的完整 Q&A。', bold=True, sz=10, clr='C00000', sa=8)

SEC3('核心架构')

Q('Q1：请介绍 HiSi DevTool 的整体架构？')
CODE('''Frontend (Vue 3 + TypeScript + Element Plus)
    |  REST API + WebSocket (Claude Terminal)
    |  SSE (Merge Analysis 实时推送)
Backend (Spring Boot 3 / JDK 21)
    +-- 代码解析层: JavaParser (Java AST) + ANTLR4 (Python AST)
    +-- 图谱层: Neo4j 5.11+ (Graph + Native Vector Index)
    +-- 检索层: 9 种查询策略 + RRF (k=60) 融合排序
    +-- AI 层: 智谱AI (GLM-4-flash + embedding-3) / Claude API
    +-- 分析层: APM Trace / RAM 需求分析 / 合并影响分析
    +-- 终端层: PTY (Claude CLI 进程管理)
    +-- 存储层: SQLite (元数据/会话) + Caffeine (本地缓存)
    +-- 调度层: Cron (定时图谱刷新)''')

SEC3('知识图谱构建')

Q('Q2：代码知识图谱是怎么构建的？')
ANS('图谱构建分为四步：')
B('1. 代码扫描：遍历项目目录，识别 Java 和 Python 源文件。')
B('2. AST 解析：Java 代码用 JavaParser 解析（内置 Java 语法支持），Python 代码用 ANTLR4 + 自定义 Python 3.8-3.12 语法文件解析。Visitor 模式遍历 AST，提取类、方法、调用关系、注解、SQL 语句等信息。')
B('3. 向量化：对每个方法生成三种 embedding——descriptionEmbedding（方法的自然语言描述）、codeEmbedding（方法代码片段）、sqlEmbedding（方法中的 SQL 语句）。使用智谱 AI embedding-3 模型，维度 2048。')
B('4. 图谱写入：将提取的节点和关系批量写入 Neo4j，使用 MERGE 语句保证幂等（重复扫描不会产生重复数据）。')
ANS('节点类型：')
TBL(['节点类型','关键字段','说明'],[
    ['MethodNode','className, methodName, signature, 3种embedding, complexity, projectPath, language','核心节点，每个方法一个'],
    ['EntryPointNode','uri, httpMethod, type(CONTROLLER/SCHEDULED/MQ/FEIGN), className, methodName','API 入口点'],
    ['ServiceNode','className, projectPath, publicProjectPath','类/服务节点'],
    ['SqlNode','sqlText, type(SELECT/INSERT/UPDATE/DELETE), projectPath','SQL 语句节点'],
    ['DataModelNode','name, className, projectPath','数据模型节点'],
    ['GenerationCheckpointNode','projectPath, timestamp, status','图谱生成检查点（增量更新用）'],
])
ANS('关系类型：')
TBL(['关系','含义','说明'],[
    ['CALLS','方法调用方法','核心关系，构成调用图'],
    ['EXTENDS','类继承','Java extends 关系'],
    ['IMPLEMENTS','接口实现','Java implements 关系'],
    ['HAS_SQL','方法包含 SQL','方法内嵌的 SQL 语句'],
    ['EXPOSES','入口点暴露方法','Controller -> Service 方法'],
    ['USES_MODEL','方法使用数据模型','方法涉及的实体/DTO'],
])

A('追问 1：如何解析 Java 代码？')
ANS('使用 JavaParser 库（3.25+ 版本）进行 Java 代码的静态分析：')
ABL('JavaParser.parse(path) 将 .java 文件解析为 AST（抽象语法树），返回 CompilationUnit 对象。')
ABL('实现 VoidVisitorAdapter 子类，重写 visit() 方法遍历各类 AST 节点：visit(ClassOrInterfaceDeclaration) 提取类定义和继承关系；visit(MethodDeclaration) 提取方法签名、参数、返回类型；visit(MethodCallExpr) 提取方法调用关系；visit(SingleMemberAnnotation/NormalAnnotation) 提取注解信息（如 @RequestMapping、@Scheduled）。')
ABL('对于 MyBatis Mapper 接口：解析 @Select/@Insert/@Update/@Delete 注解中的 SQL 语句，提取 SQL 类型和表名。')
ABL('对于 Controller：从 @RequestMapping/@GetMapping/@PostMapping 注解中提取 URI 和 HTTP 方法，创建 EntryPointNode。')

A('追问 2：如何解析 Python 代码？')
ANS('Python 是动态语言，没有类型注解的强制约束，解析比 Java 困难得多：')
ABL('使用 ANTLR4 + Python 3.8-3.12 的官方语法文件（.g4），生成 Lexer 和 Parser。ANTLR4 的优势是支持多种 Python 版本的语法（包括 walrus operator :=、match/case 等新语法）。')
ABL('实现 Visitor 模式遍历 AST：visit_FunctionDef 提取函数定义（名称、参数、装饰器、返回类型注解）；visit_ClassDef 提取类定义（继承关系、方法列表）；visit_Call 提取函数调用关系；visit_Import/visit_ImportFrom 提取导入关系。')
ABL('限制：(1) 动态类型导致无法完全确定变量类型，只能从 type hints 和赋值推断；(2) 装饰器（如 @app.route）需要特殊处理才能识别 FastAPI/Django 的路由入口点；(3) eval/exec 动态执行无法静态分析。')

A('追问 3：图谱如何更新？增量还是全量？')
ANS('采用增量更新策略，通过 Checkpoint 机制实现：')
ABL('首次扫描：全量扫描项目所有源文件，生成完整图谱。记录 checkpoint：GenerationCheckpointNode（projectPath、timestamp、status=COMPLETED）。')
ABL('增量扫描：对比当前文件的修改时间与上次 checkpoint 时间，只扫描变更的文件（新增/修改/删除）。对变更文件重新解析，更新对应的节点和关系。')
ABL('定时任务：通过 SQLite 的 kg_schedule 表配置 cron 表达式，定时触发增量扫描（如每 30 分钟一次）。')
ABL('手动触发：用户在界面上点击「刷新图谱」触发全量或增量扫描。')

SEC3('混合检索引擎')

Q('Q3：9 种查询策略分别是什么？如何用 RRF 融合？')
ANS('9 种查询策略分为三类：语义检索、关键词检索、图遍历检索。')
TBL(['类别','策略','原理','优势'],[
    ['语义','SEMANTIC_DESCRIPTION','descriptionEmbedding 向量相似度搜索','理解自然语言描述的语义'],
    ['语义','SEMANTIC_CODE','codeEmbedding 向量相似度搜索','理解代码实现的语义'],
    ['语义','SEMANTIC_SQL','sqlEmbedding 向量相似度搜索','理解 SQL 语义'],
    ['关键词','KEYWORD_METHOD_NAME','方法名 Lucene 全文搜索','精确匹配方法名'],
    ['关键词','KEYWORD_CLASS_NAME','类名 Lucene 全文搜索','精确匹配类名'],
    ['关键词','KEYWORD_SQL_TEXT','SQL 文本 Lucene 全文搜索','精确匹配 SQL 片段'],
    ['图遍历','GRAPH_UPSTREAM','从目标方法向上游遍历 CALLS 关系','找到调用该方法的上游链路'],
    ['图遍历','GRAPH_DOWNSTREAM','从目标方法向下游遍历 CALLS 关系','找到该方法调用的下游链路'],
    ['图遍历','GRAPH_ENTRY_POINT','从 EntryPointNode 反向遍历到目标方法','找到通过哪些 API 可以触达该方法'],
])

ANS('RRF 融合排序原理：')
ANS('Reciprocal Rank Fusion 是一种无需分数归一化的多路结果融合算法。核心公式：')
CODE('RRF_score(d) = SUM( 1 / (k + rank_i(d)) )  其中 k=60')
ANS('对每种策略的排名结果，按 1/(k+rank) 计算分数并求和。k=60 是论文推荐值，用来平滑排名差异。例如：一个方法在语义搜索中排第 1、关键词搜索排第 5、图遍历排第 3，其 RRF 分数 = 1/(60+1) + 1/(60+5) + 1/(60+3) = 0.0164 + 0.0154 + 0.0159 = 0.0477。')
ANS('为什么选择 RRF 而不是加权平均？(1) 不同策略的分数量纲不同（向量余弦相似度 0-1，关键词 TF-IDF 分数范围不确定），直接加权不公平；(2) RRF 只依赖排名，不需要归一化；(3) 实验证明 RRF 在混合检索场景中效果优于简单加权（参考论文：Cormack et al., 2009）。')

SEC3('APM Trace 分析')

Q('Q4：APM 模块的数据流？如何实现精准插桩？')
CODE('''数据流:
应用代码 -> OTel Agent (自动/手动埋点)
         -> OTLP/HTTP Protobuf 格式
         -> HiSi DevTool /api/otlp/traces 接收
         -> Protobuf 反序列化 (opentelemetry-proto SDK)
         -> KG 精准插桩 (只保留关键 span)
         -> silent_catch 检测 (代码静态分析)
         -> LLM 根因诊断 (Claude API)
         -> 诊断报告返回前端''')

ANS('精准插桩的核心思想：不是全量采集所有 span（会产生海量数据），而是根据知识图谱中的信息选择性采集关键 span：')
ABL('KG 知道项目的完整调用图（所有 Controller -> Service -> DAO 的调用链）。')
ABL('根据图谱中的入口点（Controller/定时任务/MQ 监听器）确定关键路径。')
ABL('只采集关键路径上的 span（入口点 -> 核心服务方法 -> 数据库操作），忽略框架内部的 span（如 Spring AOP 代理、Filter 链等）。')
ABL('效果：数据量减少 90%+，同时保留了最有诊断价值的链路信息。')

ANS('silent_catch 检测：')
ABL('问题：代码中存在 catch 块只打了日志但没有抛出异常（或 catch 块是空的），这会吞掉异常，导致上层调用方无法感知错误，问题难以定位。')
ABL('检测方式：结合两种方法——(1) 静态分析：用 JavaParser 扫描所有 catch 块，检查是否有 throw/重新抛出；(2) 运行时分析：对比 trace 中正常结束的 span 和代码中包含 try-catch 的方法，找出「应该有异常但没有抛出」的可疑情况。')
ABL('LLM 诊断：将可疑的 silent_catch 代码上下文 + 运行时 trace 信息传给 Claude API，由 LLM 分析异常根因并给出修复建议（如「catch 块应该重新抛出或返回错误码」）。')

A('追问：OTLP Protobuf 如何解析？')
ANS('OTLP（OpenTelemetry Protocol）是 OpenTelemetry 的标准数据传输协议，使用 Protobuf 编码：')
ABL('引入依赖：io.opentelemetry.proto:opentelemetry-proto:1.x。')
ABL('接收端点：Spring Boot Controller 接收 HTTP POST 请求，Content-Type 为 application/x-protobuf。')
ABL('解析：TracesData.parseFrom(inputStream) 将 Protobuf 二进制数据反序列化为 Java 对象。')
ABL('遍历：TracesData -> ResourceSpans（一个应用的所有 span）-> ScopeSpans（一个 instrumentation scope 的 span）-> Span（单个操作）。每个 Span 包含 traceId、spanId、parentSpanId、name、startTimeUnixNano、endTimeUnixNano、status、attributes 等信息。')
ABL('通过 parentSpanId 重建 span 之间的父子关系，构建完整的调用树。')

SEC3('RAM 需求分析引擎')

Q('Q5：RAM 需求分析引擎的 5 个节点分别做什么？')
CODE('''[澄清 Clarify] -> [影响 Impact] -> [代码实现 Implement] -> [验证 Verify] -> [技术方案 TechPlan]''')
ANS('每个节点的详细职责：')
ABL('Clarify（澄清）：接收用户的模糊需求描述（如「优化报表查询性能」），通过 LLM 分析需求中缺失的信息（哪个报表？当前性能指标？目标指标？），生成结构化的需求澄清问题列表。用户回答后，生成标准化的需求文档。')
ABL('Impact（影响分析）：将需求中的关键词（如「报表查询」）映射到知识图谱中的节点，通过图遍历定位受影响的代码范围。例如：找到所有报表相关的 Controller -> Service -> DAO -> SQL，以及上下游依赖。输出受影响的文件/方法清单和影响范围评估。')
ABL('Implement（代码实现）：将需求描述 + 影响分析结果 + 代码上下文一起传给 LLM，生成代码变更方案。包括：哪些文件需要修改、具体修改什么、新增什么代码。LLM 生成的方案需要结合 KG 的调用关系保证修改的完整性（不遗漏受影响的代码）。')
ABL('Verify（验证）：自动验证变更是否满足需求——(1) 生成测试建议（基于影响分析确定回归范围）；(2) 检查代码变更是否完整覆盖了需求中提到的所有场景；(3) 输出验证清单供人工确认。')
ABL('TechPlan（技术方案）：汇总前面四个节点的输出，生成完整的技术方案文档——需求描述、影响范围、变更方案、测试计划、风险评估。输出为 Markdown 格式，可直接用于技术评审。')

A('追问 1：事件溯源 + InputsHasher 最小重算如何实现？')
ANS('事件溯源（Event Sourcing）：')
ABL('每个节点的输入和输出都作为事件记录到 SQLite 的 agent_event 表中。字段包括：session_id（属于哪个分析会话）、seq（序列号）、type（事件类型如 CLARIFY_INPUT/CLARIFY_OUTPUT）、payload（JSON 格式的事件数据）、timestamp。')
ABL('完整的执行过程可以通过重放事件序列还原，支持回溯和调试。')
ANS('InputsHasher 最小重算：')
ABL('每个节点执行前，对其输入（包括上游节点的输出、用户的原始需求、KG 查询结果等）计算 SHA-256 hash。')
ABL('将 hash 与上次执行时的 hash 对比。如果 hash 相同，说明输入没有变化，直接跳过该节点，复用上次的输出结果。')
ABL('效果：如果用户只修改了 Clarify 阶段的需求描述，只有 Clarify 及其下游节点会重新执行，Impact 和前面的步骤直接复用。在 DAG 有 5 个节点的场景下，平均可以节省 60% 的 LLM 调用成本和执行时间。')

A('追问 2：HITL（Human-In-The-Loop）如何实现？')
ANS('HITL 机制允许用户在关键节点暂停并介入：')
ABL('节点执行到关键决策点时（如 Clarify 需要用户回答问题、Implement 需要用户确认方案），将状态保存为 PENDING_HUMAN_REVIEW，暂停 DAG 执行。')
ABL('前端通过 WebSocket 或轮询获取节点状态，展示 PENDING 节点的输出供用户审阅。')
ABL('用户可以：(1) 确认通过，继续执行下游节点；(2) 修改内容后确认，使用修改后的输入继续；(3) 驳回并重新执行当前节点。')
ABL('所有人工修改都记录在事件日志中，保证可追溯。')

SEC3('合并影响分析')

Q('Q6：合并影响分析模块如何工作？')
ANS('完整的数据流：')
B('1. 触发：用户在界面上指定目标分支（如 main）和源分支（如 feature/xxx），触发合并分析。')
B('2. JGit Diff：使用 JGit 库（org.eclipse.jgit）计算两个分支的差异——ModifiedFiles（修改的文件列表）、AddedLines/DeletedLines（新增/删除的行）、DiffEntry（每个文件的 diff 详情）。')
B('3. 变更映射到 KG：将 diff 中涉及的文件路径和方法名映射到 KG 中的 MethodNode。例如：修改了 TaxCalculator.java 的 calculateVAT 方法，找到对应的 MethodNode。')
B('4. 三层上游追溯：')
B('   第一层：直接调用方 — 查询 CALLS 关系的反向，找到谁直接调用了 calculateVAT（如 Controller 的 submitTaxReturn 方法）。')
B('   第二层：调用方的调用方 — 继续向上遍历，找到 submitTaxReturn 的调用方（如前端 API /api/tax/submit）。')
B('   第三层：入口点 — 遍历到 EntryPointNode，确定哪些 API/定时任务/MQ 监听器最终会受影响。')
B('5. LLM 测试建议：将三层追溯的结果（受影响的入口点、调用链路、变更代码的上下文）传给 Claude API，生成测试范围建议——哪些功能需要回归测试、重点测试什么场景、可能的风险点。')
B('6. SSE 推送：整个分析过程中，每个步骤完成后通过 SSE（Server-Sent Events）向前端推送进度（如「正在解析 Diff...」「正在追溯上游...」），用户可以实时看到分析进展。')

A('追问：为什么用 SSE 而不是 WebSocket？')
ANS('SSE vs WebSocket 的选择考量：')
ABL('数据流向：合并分析是典型的单向推送场景（服务器 -> 客户端），不需要客户端向服务器频繁发送数据。SSE 是 HTTP 单向推送（server -> client），WebSocket 是双向通信。SSE 更简单直接。')
ABL('协议简单度：SSE 基于标准 HTTP，不需要额外的协议升级（WebSocket 需要 Upgrade: websocket 握手），防火墙和代理兼容性更好。')
ABL('自动重连：SSE 内置自动重连机制（浏览器 EventSource API 自动处理），WebSocket 需要手动实现重连逻辑。')
ABL('HTTP/2 多路复用：在 HTTP/2 下，多个 SSE 流可以复用同一个 TCP 连接，不会阻塞其他请求。')
ANS('什么时候用 WebSocket：(1) 需要双向通信（如 Claude Terminal，用户输入命令 -> 服务器返回输出）；(2) 需要传输二进制数据；(3) 需要极低延迟的实时通信。')

SEC3('LLM 集成')

Q('Q7：智谱 AI 和 Claude API 如何集成？')
ANS('统一接口设计：所有 LLM 调用封装在 LlmService 中，对外暴露统一的 chat(messages) -> response 接口，内部根据场景路由到不同的 LLM 提供商。')
ABL('智谱 AI：使用 GLM-4-flash 模型进行文本生成（需求分析、代码审查），embedding-3 模型进行文本向量化（2048 维）。调用方式：OpenAI 兼容的 /v1/chat/completions 接口。')
ABL('Claude API：通过 dmxapi 代理服务调用（因为直接调用 Anthropic API 存在网络限制）。用于 APM 根因诊断等需要更强推理能力的场景。Claude 的推理能力在复杂代码分析场景中优于 GLM-4-flash。')
ABL('隔离通道：不同 LLM 用于不同场景，互不影响。智谱 AI 不可用时不影响 Claude API 的使用，反之亦然。')

A('追问 1：如何处理 LLM 的幻觉？')
ANS('LLM 幻觉是指模型生成看似合理但实际上是错误的内容。在代码分析场景中，幻觉可能导致错误的代码建议。我们的对策：')
ABL('KG 事实约束：在 Prompt 中明确告知 LLM「只基于提供的代码上下文回答，不要推测不存在的代码」。将 KG 中的真实代码结构（类名、方法名、调用关系）作为上下文传入。')
ABL('输出校验：LLM 返回的结果（如代码建议）会与 KG 中的节点进行交叉验证。如果 LLM 建议调用一个 KG 中不存在的方法，标记为可疑。')
ABL('结构化输出：使用 JSON Schema 约束 LLM 的输出格式，减少自由文本带来的不确定性。')
ABL('人工审核：关键决策点（如 RAM 的 Implement 阶段）通过 HITL 机制让用户审核 LLM 的输出。')

A('追问 2：Token 成本如何控制？')
ANS('LLM API 调用按 Token 计费，大规模使用时成本可观。我们的成本控制策略：')
ABL('Prompt 精简：精心设计 Prompt，只传入必要的上下文（如只传入相关方法的代码，而不是整个文件）。使用 KG 精准定位代码范围，避免传入无关代码。')
ABL('模型分级：简单任务（如文本摘要）用轻量模型 GLM-4-flash，复杂推理任务（如根因诊断）用 Claude。不是所有任务都需要最强的模型。')
ABL('缓存：对相似查询的 embedding 结果缓存（Neo4j 的向量索引本身就是持久化的），避免重复调用 embedding API。对 LLM 的结果也做缓存（相同输入+相同模型+相同参数 = 相同输出）。')
ABL('InputsHasher：RAM 引擎的最小重算机制，跳过输入未变化的节点，直接复用上次的 LLM 结果。')
ABL('批量调用：embedding 生成支持批量接口（一次传入多个文本），减少 API 调用次数。')

print('Part 3 done')
doc.save(OUTPUT)
print(f'Saved {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')
