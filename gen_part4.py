#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Part 4: System Design + Behavioral + Algorithm + Appendix + 反问面试官"""
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'C:\Users\47583\projects\hisi_dev_tool v5.0\沈江春_面试准备手册_完整版.docx'
doc = Document(OUTPUT)

# ── helpers ──
def P(text, bold=False, sz=None, clr=None, align=None, sa=None, sn=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if sz: r.font.size = Pt(sz)
    if clr: r.font.color.rgb = RGBColor(*clr)
    if align: p.alignment = align
    pf = p.paragraph_format
    if sa is not None: pf.space_after = Pt(sa)
    if sn is not None: pf.space_before = Pt(sn)
    return p

def B(text, lv=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if lv > 0:
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25 * lv)
    return p

def TBL(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = h
        for r2 in c.paragraphs[0].runs:
            r2.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = str(val)
    return tbl

def CODE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = 'Consolas'
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return p

def A(title):
    P(title, bold=True, sz=11, clr=(0, 100, 180), sa=2)

def ANS(text):
    P(text, sz=10.5, sa=2)

def ABL(text):
    B(text)

def Q(text):
    P(text, bold=True, sz=12, sa=2)

def SEC1(t):
    P(t, bold=True, sz=16, clr=(0, 51, 102), sa=6, sn=12)
def SEC2(t):
    P(t, bold=True, sz=14, clr=(0, 80, 140), sa=4, sn=8)
def SEC3(t):
    P(t, bold=True, sz=12, clr=(0, 100, 160), sa=3, sn=6)

# ════════════════════════════════════════
# SECTION 6: SYSTEM DESIGN
# ════════════════════════════════════════
SEC1('六、系统设计')
P('系统设计面试考察的是结构化思维和权衡能力。以下是结合简历项目的高频系统设计题。', sz=10.5, sa=4)

# 6.1 多租户 SaaS
SEC2('6.1 多租户 SaaS 架构设计')
Q('请设计一个支持 300+ 租户的纳税申报 SaaS 系统')
ANS('需求分析：')
ABL('功能：多租户独立数据隔离、税务申报、发票管理、报表统计')
ABL('非功能：300+ 租户、每租户 100-1000 用户、SLA 99.9%、数据严格隔离')
ABL('约束：财务数据必须准确（分）、不能跨租户泄漏')

ANS('整体架构（四层）：')
TBL(['层次', '技术选型', '职责'], [
    ['接入层', 'Nginx + Spring Cloud Gateway', '路由、限流、租户识别（Header/X-Tenant-ID）、JWT 验证'],
    ['应用层', 'Spring Boot 微服务 × 4', '申报服务、发票服务、报表服务、用户服务'],
    ['数据层', 'PostgreSQL Schema 隔离 + Redis', '每租户独立 Schema，缓存热点配置'],
    ['基础设施', 'K8s + Kafka + MinIO', '容器编排、异步消息、文件存储'],
])

ANS('多租户隔离方案选型：')
TBL(['方案', '隔离级别', '优点', '缺点', '适用场景'], [
    ['独立数据库', '最高', '完全隔离，可独立备份恢复', '成本最高，300+ 实例管理复杂', '金融/政务'],
    ['共享数据库独立 Schema', '中高', 'Schema 级隔离，管理较简单', 'Schema 数量有上限（PG 约几千）', '中大型 SaaS（我们的选择）'],
    ['共享 Schema + tenant_id 列', '中', '成本最低，资源利用率最高', '隔离最弱，SQL 注入风险高', '小型/免费 SaaS'],
])
ANS('我们选择方案 2（独立 Schema），核心理由：')
ABL('财务数据敏感性要求 Schema 级隔离，比行级隔离更安全')
ABL('300 租户在 PostgreSQL 的 Schema 上限内')
ABL('可以对不同租户做独立的备份、恢复、vacuum')
ABL('通过 Flyway 管理 Schema 迁移，统一版本')

ANS('关键设计细节 —— 租户上下文传递：')
CODE('''// 1. Gateway 层提取租户 ID
@Component
public class TenantFilter implements GlobalFilter {
    @Override
    public Mono<Void> filter(ServerWebExchange exchange, GatewayFilterChain chain) {
        String tenantId = exchange.getRequest().getHeaders().getFirst("X-Tenant-ID");
        // 验证 tenantId 合法性（JWT claims 中也包含 tenantId，两者必须一致）
        return chain.filter(exchange);
    }
}

// 2. 应用层 TenantContext（ThreadLocal）
public class TenantContext {
    private static final ThreadLocal<String> CURRENT = new ThreadLocal<>();
    public static void set(String tenantId) { CURRENT.set(tenantId); }
    public static String get() { return CURRENT.get(); }
    public static void clear() { CURRENT.remove(); }
}

// 3. MyBatis 拦截器动态切换 Schema
@Intercepts({@Signature(type=StatementHandler.class, method="prepare", args={Connection.class,Integer.class})})
public class TenantSchemaInterceptor implements Interceptor {
    @Override
    public Object intercept(Invocation invocation) throws Throwable {
        String schema = "tenant_" + TenantContext.get();
        Connection conn = (Connection) invocation.getArgs()[0];
        conn.setSchema(schema);
        return invocation.proceed();
    }
}''')

ANS('数据库连接池优化：')
ABL('HikariCP 配置：每个 Schema 独立连接池不现实（300 × 10 = 3000 连接），采用共享连接池 + 动态切换 Schema')
ABL('连接池大小：maximumPoolSize = CPU 核数 × 2 + 磁盘数（约 20），配合 PgBouncer 做连接复用')
ABL('设置 search_path 而非 setSchema()，避免连接归还后 Schema 残留')

ANS('跨租户报表需求处理：')
ABL('管理后台需要跨租户统计（如总申报数、总税额），直接查各 Schema 不现实')
ABL('方案：Kafka 消费各租户的申报完成事件 → 写入公共 Schema 的聚合表 aggregate_report')
ABL('聚合表只存脱敏后的统计数据，不含具体财务明细')

A('追问1：如何防止跨租户数据泄漏？')
ANS('多层防护：')
ABL('代码层：所有 SQL 查询必须带 tenant_id 条件，通过 MyBatis 拦截器自动注入 WHERE tenant_id = ? 条件')
ABL('Schema 层：每个请求开始时设置 search_path，SQL 执行在租户 Schema 内，天然隔离')
ABL('API 层：Spring Security + 自定义 TenantPermissionEvaluator，验证当前用户是否有权访问目标租户')
ABL('审计层：所有数据访问日志记录 tenant_id，定期审计异常访问模式（如 A 租户用户频繁查询 B 租户）')
ABL('测试层：自动化测试中增加跨租户泄漏测试 case，验证不同租户 token 无法访问其他租户数据')

A('追问2：租户 Schema 迁移如何做？')
ANS('使用 Flyway 做统一迁移：')
ABL('编写 migration SQL 脚本，版本号如 V2024.01__add_invoice_table.sql')
ABL('部署时遍历所有租户 Schema，依次执行迁移：for schema in tenant_*: flyway migrate -schemas=$schema')
ABL('支持灰度迁移：先对 5 个试点租户执行，观察无异常后再全量执行')
ABL('迁移前自动备份目标 Schema（pg_dump -n $schema），失败可回滚')
ABL('监控迁移进度和耗时，Schema 超过 1000 时考虑分批执行（每批 50 个，间隔 5 分钟）')

# 6.2 高性能搜索
SEC2('6.2 高性能代码搜索系统设计')
Q('设计一个支持向量 + 关键词混合搜索的代码知识检索系统')
ANS('需求分析：')
ABL('数据规模：100 万+ 代码方法节点，500 万+ 调用关系边')
ABL('查询模式：语义搜索（自然语言描述找代码）、精确搜索（类名+方法名）、调用链查询（上下游 N 层）')
ABL('性能要求：单次搜索 < 500ms（P99），支持 50 并发')
ABL('一致性：准实时（分钟级延迟可接受）')

ANS('整体架构：')
TBL(['组件', '职责', '技术选型'], [
    ['数据采集层', '源码解析、AST 提取、向量化', 'ANTLR4 + Tree-sitter + Sentence Transformers'],
    ['存储层', '结构化存储 + 向量索引 + 图关系', 'PostgreSQL + pgvector + Apache AGE'],
    ['查询层', '混合检索 + 图遍历 + 结果融合', 'Spring Boot REST API'],
    ['缓存层', '热点查询缓存', 'Redis + Caffeine L1'],
])

ANS('索引构建流程：')
CODE('''1. 源码扫描 → ANTLR4 解析 → 生成 AST
2. AST Visitor 遍历 → 提取 Class/Method/Field 节点 → 写入 PostgreSQL
3. 调用关系提取 → 静态分析 + 反射解析 → 写入 Apache AGE 图
4. 方法节点文本化 → "ClassName.methodName: docstring + 参数签名"
5. Sentence Transformer 编码 → 768 维向量 → 存入 pgvector
6. 向量索引：IVFFlat (lists=100) 或 HNSW (m=16, ef_construction=200)''')

ANS('查询执行流程（RRF 混合检索）：')
CODE('''用户输入: "处理支付回调的方法"

Step 1: 关键词过滤（精确缩小范围）
  → PostgreSQL WHERE code_text ILIKE '%支付%' OR '%callback%' OR '%payment%'
  → 候选集 C_keyword (约 500-5000 个节点)

Step 2: 向量语义搜索（在候选集内）
  → pgvector: SELECT * FROM methods ORDER BY embedding <=> $query_vec LIMIT 100
  → 返回 C_vector (100 个最相似节点)

Step 3: 图遍历扩展（补全调用链上下文）
  → 对 C_vector 中 top-10 节点做 BFS/DFS，获取 depth=2 的上下游
  → 扩展集 C_graph

Step 4: RRF 融合排序
  → RRF_score(d) = Σ 1/(k + rank_i(d))，k=60
  → 合并三个候选集，按 RRF 分数降序，返回 top-20''')

ANS('性能优化策略：')
ABL('向量索引：HNSW 替代 IVFFlat，查询速度快 3-5 倍，内存占用大但 100 万节点可接受（约 3GB）')
ABL('分页缓存：相同查询参数的前 3 页结果缓存到 Redis，TTL 5 分钟')
ABL('异步图遍历：图遍历是最耗时的步骤（200-500ms），使用 CompletableFuture 异步执行，与向量搜索并行')
ABL('连接池调优：pgvector 查询是 CPU 密集型，HikariCP 连接数 = CPU 核数')

A('追问1：为什么选 pgvector 而非 Milvus/Weaviate？')
ANS('技术选型权衡：')
ABL('运维成本：pgvector 是 PostgreSQL 扩展，无需额外组件；Milvus 需要 etcd + MinIO + 多个组件')
ABL('数据一致性：pgvector 和业务数据在同一个 PostgreSQL 实例中，天然事务一致；独立向量库需要额外同步机制')
ABL('规模匹配：100 万 × 768 维 ≈ 3GB，pgvector + HNSW 完全可以处理；超过 1000 万才需要考虑 Milvus')
ABL('开发效率：直接用 SQL 查询，不需要学习新 API，团队上手快')
ABL('劣势：pgvector 在超大规模（1亿+）和高吞吐场景下性能不如 Milvus，但当前规模足够')

A('追问2：如何保证向量索引与数据库的一致性？')
ANS('同步方案：')
ABL('源码变更 → Git Webhook → 触发增量解析任务 → 更新 PostgreSQL 节点 + 重新生成向量 → 更新 pgvector')
ABL('事务保证：方法记录更新和向量更新在同一个 PostgreSQL 事务中，不会出现"方法已更新但向量还是旧的"')
ABL('全量重建：每天凌晨 2 点跑一次全量对比（源码 hash vs 数据库 hash），修复增量遗漏的节点')
ABL('版本控制：向量表增加 source_hash 字段，比对源码变更决定是否需要重新编码')

# 6.3 API 限流
SEC2('6.3 分布式限流系统设计')
Q('设计一个支持多维度限流的 API 网关')
ANS('需求分析：')
ABL('限流维度：全局限流（总 QPS）、用户级限流（每用户 QPS）、接口级限流（特定接口 QPS）')
ABL('算法：令牌桶（允许突发）或滑动窗口（精确控制）')
ABL('分布式：多实例部署，限流状态需要共享')
ABL('性能：限流判断 < 1ms，不能成为瓶颈')

ANS('技术方案：Redis + Lua 脚本实现滑动窗口限流')
CODE('''-- Redis Lua 滑动窗口限流脚本
local key = KEYS[1]           -- 限流 key，如 "rate:user:123"
local window = tonumber(ARGV[1])  -- 窗口大小（秒）
local limit = tonumber(ARGV[2])   -- 窗口内允许的最大请求数
local now = tonumber(ARGV[3])     -- 当前时间戳（毫秒）

-- 删除窗口外的旧数据
redis.call('ZREMRANGEBYSCORE', key, 0, now - window * 1000)
-- 统计窗口内的请求数
local count = redis.call('ZCARD', key)
if count < limit then
    -- 未超限，加入当前请求
    redis.call('ZADD', key, now, now .. ':' .. math.random(100000))
    redis.call('EXPIRE', key, window)
    return 1  -- 允许
else
    return 0  -- 拒绝
end''')

ANS('多维度限流实现：')
CODE('''@Component
public class RateLimiterService {
    @Autowired
    private StringRedisTemplate redis;

    public boolean isAllowed(String userId, String apiPath) {
        // 维度1: 全局限流 — 10000 QPS
        boolean globalOk = checkLimit("rate:global", 1, 10000);
        // 维度2: 用户级 — 100 QPS
        boolean userOk = checkLimit("rate:user:" + userId, 1, 100);
        // 维度3: 接口级 — 特定接口 500 QPS
        boolean apiOk = checkLimit("rate:api:" + apiPath, 1, 500);

        return globalOk && userOk && apiOk;
    }

    private boolean checkLimit(String key, int windowSec, int limit) {
        Long result = redis.execute(script,
            List.of(key),
            String.valueOf(windowSec), String.valueOf(limit),
            String.valueOf(System.currentTimeMillis()));
        return result != null && result == 1L;
    }
}''')

ANS('限流降级策略：')
ABL('Redis 不可用时降级为本地令牌桶限流（Guava RateLimiter），防止无限放行')
ABL('返回 HTTP 429 + Retry-After Header，客户端可自动重试')
ABL('限流日志记录到 Kafka，用于后续分析流量模式和调整限流阈值')

A('追问1：令牌桶 vs 滑动窗口 vs 漏桶的区别？')
ANS('三种算法对比：')
TBL(['算法', '原理', '突发处理', '适用场景'], [
    ['令牌桶', '按固定速率放令牌，请求消耗令牌', '允许突发（桶内有余量时）', 'API 限流（推荐）'],
    ['滑动窗口', '统计窗口内请求数，超限拒绝', '无突发，精确控制', '严格限流场景'],
    ['漏桶', '请求入桶，固定速率流出', '无突发，平滑流量', '流量整形/削峰'],
])
ANS('我们选择滑动窗口的原因：')
ABL('Redis + Lua 实现简单且原子性有保障')
ABL('不需要突发处理能力（API 限流更看重精确控制）')
ABL('可以通过调整窗口大小（1s/1min/1h）实现不同粒度的限流')

A('追问2：分布式限流如何保证原子性？')
ANS('核心是 Lua 脚本在 Redis 中原子执行：')
ABL('Redis 执行 Lua 脚本是单线程原子操作，不会出现并发下"读-判断-写"的竞态条件')
ABL('ZCARD + ZADD 在同一个 Lua 脚本中，中间不会被其他请求插入')
ABL('多个应用实例共享同一个 Redis，天然分布式一致性')
ABL('如果用 Redis Cluster，相同 key 会路由到同一节点，Lua 脚本原子性仍然保证')


# ════════════════════════════════════════
# SECTION 7: BEHAVIORAL INTERVIEW (STAR)
# ════════════════════════════════════════
SEC1('七、行为面试 STAR 话术')
P('行为面试用 STAR 法则（Situation-Task-Action-Result）组织回答，每个回答控制在 2-3 分钟。', sz=10.5, sa=4)

SEC2('7.1 你遇到过最有挑战的技术难题是什么？')
ANS('Situation（背景）：')
ABL('在纳税申报 SaaS 项目中，系统上线 3 个月后，某次税务申报高峰期（月末最后一天），多个大租户同时提交年度汇总报表，系统响应时间从正常的 200ms 飙升到 15 秒以上，部分请求超时，值班群收到大量用户投诉。')
ANS('Task（任务）：')
ABL('作为后端负责人，需要在 4 小时内定位问题根因并修复，因为月末是法定申报截止日，系统不可用会导致租户错过申报期限。')
ANS('Action（行动）：')
ABL('第一步（30 分钟）：快速止血 —— 通过 Grafana 定位到瓶颈在报表查询接口，临时将该接口的并发限制从 100 降到 20，用排队机制缓解压力，确保已有请求能正常完成。')
ABL('第二步（1 小时）：根因定位 —— 分析慢查询日志，发现报表 SQL 使用了 EXTRACT(YEAR FROM report_date) 函数导致全表扫描；同时 pg_stat_activity 发现大量 WAIT EVENT 是 Lock，原因是多个报表请求竞争同一个汇总临时表的写锁。')
ABL('第三步（2 小时）：修复 —— (1) 将函数条件改写为范围查询，创建复合索引，查询时间从 10s+ 降到 150ms；(2) 将汇总临时表改为独立的物化视图（MATERIALIZED VIEW），消除写锁竞争；(3) 增加报表结果缓存（Caffeine，30 分钟 TTL），重复查询直接命中缓存。')
ANS('Result（结果）：')
ABL('修复后系统恢复正常，同日晚高峰 P99 响应时间 300ms，无超时。后续一个月内未再出现类似问题。该优化方案沉淀为团队的"报表查询优化 Checklist"，后续其他报表接口也按此模式优化。')

A('追问：如果时间更充裕，你会做什么额外优化？')
ANS('更深层的优化（当时没来得及做）：')
ABL('引入读写分离：报表查询走只读副本，与 OLTP 业务隔离')
ABL('预计算聚合：每日凌晨物化常用维度的聚合数据，查询时直接读聚合结果而非实时计算')
ABL('分级 SLA：大租户和小租户的报表查询资源池隔离，防止大租户拖垮小租户')

SEC2('7.2 描述一次你推动技术改进的经历')
ANS('Situation：')
ABL('2024 年初加入华为后，发现团队的 JDK 还停留在 8，Spring Boot 用的 2.x，很多新特性用不了（如 Records、Sealed Classes、Virtual Threads），代码中充斥着冗长的 getter/setter/toString 和手动异常处理。')
ANS('Task：')
ABL('作为项目组内最熟悉新 JDK 特性的成员，需要推动团队将核心服务从 JDK 8 + Spring Boot 2.x 升级到 JDK 21 + Spring Boot 3.x，同时不能中断业务迭代。')
ANS('Action：')
ABL('第一步：调研和方案 —— 花一周梳理依赖兼容性（javax → jakarta 命名空间迁移、Spring Security 6 变更、第三方库兼容性），输出 15 页的《JDK 21 升级评估报告》，列出风险点和应对方案。')
ABL('第二步：争取支持 —— 在周会上做 30 分钟技术分享，用基准数据说明升级收益（Virtual Threads 并发吞吐提升 30%+、ZGC 停顿 < 1ms），获得 TL 和 PM 的支持。')
ABL('第三步：分阶段执行 —— Phase 1（2 周）：先在测试环境完成 JDK 21 + Spring Boot 3 的依赖适配和编译通过；Phase 2（1 周）：灰度发布到 1 个非核心服务，观察 3 天；Phase 3（2 周）：全量升级其他服务，每升级一个服务都跑完整回归测试。')
ABL('第四步：知识沉淀 —— 编写《JDK 21 升级操作手册》和《常见问题 FAQ》，组织 2 次团队培训，确保每个人都能独立处理升级过程中的问题。')
ANS('Result：')
ABL('6 周完成全组服务升级，零线上事故。升级后 GC 停顿时间从平均 50ms 降到 < 1ms（ZGC），接口 P99 响应时间下降 15%。该方案被其他项目组参考，TL 在部门技术分享会上做了推广。')

A('追问：过程中遇到的最大阻力是什么？如何克服？')
ANS('最大阻力是第三方依赖的兼容性问题：')
ABL('我们依赖的一个内部 SDK（日志采集）还在用 javax 命名空间，无法直接在 JDK 21 + Spring Boot 3 下编译')
ABL('我花 2 天联系该 SDK 的维护团队，了解到他们计划 Q2 发布 jakarta 版本')
ABL('过渡方案：使用 Eclipse Transformer 工具在编译期自动将 javax 转换为 jakarta，解决了编译问题')
ABL('正式方案：SDK 发布 jakarta 版本后第一时间升级，移除 Transformer 依赖')
ABL('关键教训：升级前必须逐个排查所有依赖的兼容性，制定应急预案，不能假设"升级就行"')

SEC2('7.3 你如何处理线上故障？')
ANS('Situation：')
ABL('某天下午 3 点，监控告警：HiSi DevTool 的 Neo4j 查询延迟突增 10 倍，用户反馈搜索结果加载超时。')
ANS('Task：')
ABL('需要在最短时间内恢复服务可用性，同时找到根因防止复发。')
ANS('Action：')
ABL('止血（5 分钟）：检查 Neo4j 监控 → 发现 heap 使用率 95%，GC 频繁 → 临时重启 Neo4j 实例，服务恢复')
ABL('根因定位（1 小时）：分析 Neo4j query.log → 发现一条 Cypher 查询执行了全图 MATCH (n)-[*1..5]->(m)，遍历了 500 万节点 × 5 层深度 → 该查询来自新上线的"调用链深度搜索"功能，缺少最大深度限制')
ABL('修复（30 分钟）：(1) 在代码中强制限制最大深度为 3；(2) 增加查询超时 10 秒；(3) 为 Cypher 查询添加 EXPLAIN 验证索引使用情况')
ABL('复盘：输出故障复盘文档，增加"新功能上线必须跑查询性能基准测试"的 Checklist')
ANS('Result：')
ABL('服务恢复时间 5 分钟，根因修复 1.5 小时。后续增加 Neo4j 慢查询自动告警（> 2 秒），类似问题再未发生。')

A('追问：如何避免类似问题再次发生？')
ANS('建立三层防护：')
ABL('开发阶段：Cypher 查询必须包含索引提示（USING INDEX）和深度限制，Code Review 时重点检查')
ABL('测试阶段：新增"性能回归测试"，对比新版本和基准版本的查询耗时，退化超过 50% 自动阻断合并')
ABL('运维阶段：Neo4j 慢查询日志 + Grafana 告警，heap > 80% 预警，> 90% 自动扩容')

SEC2('7.4 描述一次与同事意见不一致的经历')
ANS('Situation：')
ABL('在研发管理工具链项目中，关于如何存储代码知识图谱的技术选型，我和一位同事有分歧：我倾向用 Neo4j + PostgreSQL 的方案（图数据库存调用关系，关系数据库存元数据），同事认为应该全部用 PostgreSQL + Apache AGE（PostgreSQL 的图扩展），避免引入额外组件。')
ANS('Task：')
ABL('需要在不伤害合作关系的前提下，达成技术共识，做出最优决策。')
ANS('Action：')
ABL('第一步：倾听理解 —— 我先认真听了同事的理由（运维复杂度、团队学习成本、单组件一致性），确认这些顾虑是合理的')
ABL('第二步：数据驱动 —— 我花一天做了 POC 对比测试，用真实数据集（10 万节点 + 50 万边）跑 5 种典型查询，对比 Neo4j 和 Apache AGE 的性能、内存占用和开发体验')
ABL('第三步：开放讨论 —— 把 POC 结果（包括两边的优缺点）做成表格分享给团队，让大家一起讨论，而不是"我说了算"')
ANS('Result：')
ABL('POC 数据显示：Neo4j 在多跳遍历查询（3+ 层）上快 5-10 倍，但 Apache AGE 在单跳和简单查询上性能相当，且运维成本确实更低')
ABL('最终决策：采用折中方案 —— 核心的深度调用链查询用 Neo4j，简单查询和元数据用 PostgreSQL，通过数据同步保持一致')
ABL('团队关系未受影响，同事在后续项目中主动请教我技术选型的思路')

A('追问：如果重来，你会改变什么？')
ANS('会更早做 POC 而不是口头讨论：')
ABL('当时我们先争论了一周才做 POC，浪费了时间。如果一开始就用数据说话，讨论会更高效')
ABL('也会邀请更多人参与（不只是我和同事），让决策更民主，减少个人偏好影响')

SEC2('7.5 你是如何学习新技术的？')
ANS('Situation：')
ABL('2025 年初，部门决定引入 LLM 能力到研发工具链中。当时团队大多数人对 RAG、向量搜索、Prompt Engineering 几乎没有经验。')
ANS('Task：')
ABL('作为项目组后端开发，需要在 2 个月内掌握 LLM 集成技术栈，并主导设计 RAG 管道。')
ANS('Action：')
ABL('系统学习（2 周）：阅读 LangChain 官方文档 + Andrej Karpathy 的 LLM 入门视频 + 动手跑通 3 个 RAG Demo 项目')
ABL('实践验证（2 周）：用公司内部代码库搭建了一个小型 RAG Demo，测试不同 Embedding 模型（text-embedding-ada-002 vs bge-large-zh）的效果差异')
ABL('知识输出（1 周）：写了一份《RAG 技术选型与实现指南》发到团队 Confluence，包含原理、选型对比、代码示例、踩坑记录')
ABL('落地实践（3 周）：主导设计并实现了 HiSi DevTool 的 RAG 管道（文档切分 → 向量化 → 检索 → 重排 → LLM 生成）')
ANS('Result：')
ABL('2 个月内从零到主导完成 RAG 管道设计，该文档成为团队新人入职必读材料。后续在部门技术分享上做了《LLM 在研发工具中的落地实践》的主题分享。')

A('追问：学习新技术时，你最看重什么？')
ANS('三个原则：')
ABL('最小可行：先跑通最小 Demo，再逐步增加复杂度。不要一上来就读论文/源码，先建立感性认识')
ABL('对比验证：同一件事用 2-3 种方案做 POC，用数据而非感觉判断哪个好')
ABL('输出倒逼输入：写文档/做分享是最好的学习方式，能把一件事讲清楚说明真的懂了')

SEC2('7.6 为什么离开华为？')
ANS('标准回答框架（根据实际情况调整）：')
ABL('在华为学到了很多，特别是在大型项目的工程化和团队协作方面。但我希望在 AI 工程化这个方向继续深入，目前的岗位在 AI 相关的实践机会有限。')
ABL('贵公司在这个方向上有很好的实践场景，我相信在这里能做出更有影响力的技术成果，同时也能加速自己的技术成长。')
ABL('注意：不要说华为坏话，不要提薪资，聚焦"追求成长"而非"逃离现状"。')


# ════════════════════════════════════════
# SECTION 8: ALGORITHM
# ════════════════════════════════════════
SEC1('八、算法高频题')
P('以下是面试高频算法题，按类型分类，每题给出思路 + Java 代码 + 追问。', sz=10.5, sa=4)

SEC2('8.1 LRU 缓存（手写实现）')
Q('请实现一个 LRU 缓存，支持 O(1) 的 get 和 put')
ANS('思路：HashMap + 双向链表。HashMap 存 key→Node 映射，双向链表维护访问顺序（最近访问的放头部，淘汰尾部）。')
CODE('''class LRUCache {
    private final int capacity;
    private final Map<Integer, Node> map;
    private final Node head, tail;  // 哨兵节点

    public LRUCache(int capacity) {
        this.capacity = capacity;
        this.map = new HashMap<>();
        head = new Node(0, 0);
        tail = new Node(0, 0);
        head.next = tail;
        tail.prev = head;
    }

    public int get(int key) {
        Node node = map.get(key);
        if (node == null) return -1;
        moveToHead(node);
        return node.value;
    }

    public void put(int key, int value) {
        Node node = map.get(key);
        if (node != null) {
            node.value = value;
            moveToHead(node);
        } else {
            Node newNode = new Node(key, value);
            map.put(key, newNode);
            addToHead(newNode);
            if (map.size() > capacity) {
                Node removed = removeTail();
                map.remove(removed.key);
            }
        }
    }

    private void addToHead(Node node) {
        node.next = head.next;
        node.prev = head;
        head.next.prev = node;
        head.next = node;
    }

    private void removeNode(Node node) {
        node.prev.next = node.next;
        node.next.prev = node.prev;
    }

    private void moveToHead(Node node) {
        removeNode(node);
        addToHead(node);
    }

    private Node removeTail() {
        Node node = tail.prev;
        removeNode(node);
        return node;
    }

    static class Node {
        int key, value;
        Node prev, next;
        Node(int k, int v) { key = k; value = v; }
    }
}''')

A('追问1：为什么不用 LinkedHashMap 实现？')
ANS('可以用，面试官可能要求手写，但实际项目中 LinkedHashMap 更好：')
CODE('''// JDK 自带 LinkedHashMap + accessOrder=true 即可实现 LRU
class LRUCache extends LinkedHashMap<Integer, Integer> {
    private final int capacity;

    public LRUCache(int capacity) {
        super(capacity, 0.75f, true);  // accessOrder=true
        this.capacity = capacity;
    }

    public int get(int key) {
        return super.getOrDefault(key, -1);
    }

    @Override
    protected boolean removeEldestEntry(Map.Entry<Integer, Integer> eldest) {
        return size() > capacity;
    }
}''')
ABL('但手写版本展示了对数据结构的理解，面试中更受青睐')

A('追问2：如何实现线程安全的 LRU？')
ANS('几种方案：')
ABL('方案 1：Collections.synchronizedMap() 包装，但并发性能差（全锁）')
ABL('方案 2：ConcurrentHashMap + 双向链表（自己处理并发，复杂度高）')
ABL('方案 3（推荐）：用分段锁 —— 将缓存分成 N 个 Segment，每个 Segment 独立 LRU，hash(key) % N 决定去哪个 Segment')
ABL('方案 4：Caffeine 库（生产推荐），内部使用了 Window TinyLfu 算法，性能远超简单 LRU')

SEC2('8.2 二叉树层序遍历')
Q('请实现二叉树的层序遍历，返回每一层的节点值')
CODE('''List<List<Integer>> levelOrder(TreeNode root) {
    List<List<Integer>> result = new ArrayList<>();
    if (root == null) return result;
    Queue<TreeNode> queue = new LinkedList<>();
    queue.offer(root);
    while (!queue.isEmpty()) {
        int size = queue.size();
        List<Integer> level = new ArrayList<>();
        for (int i = 0; i < size; i++) {
            TreeNode node = queue.poll();
            level.add(node.val);
            if (node.left != null) queue.offer(node.left);
            if (node.right != null) queue.offer(node.right);
        }
        result.add(level);
    }
    return result;
}''')

A('追问：如何用 DFS 实现层序遍历？')
CODE('''void dfs(TreeNode node, int depth, List<List<Integer>> result) {
    if (node == null) return;
    if (result.size() == depth) result.add(new ArrayList<>());
    result.get(depth).add(node.val);
    dfs(node.left, depth + 1, result);
    dfs(node.right, depth + 1, result);
}''')
ABL('DFS 实现更简洁，但 BFS 更直观且空间复杂度在完全二叉树时更优')

SEC2('8.3 TopK 问题')
Q('如何从 N 个元素中找到最大的 K 个？')
ANS('三种方案对比：')
TBL(['方案', '时间复杂度', '空间复杂度', '适用场景'], [
    ['排序后取前 K', 'O(N log N)', 'O(1) 或 O(N)', 'K 接近 N 时'],
    ['小顶堆', 'O(N log K)', 'O(K)', 'K 较小时（推荐）'],
    ['快速选择', 'O(N) 平均', 'O(1)', '需要原地操作'],
])

CODE('''// 方案：小顶堆
public int[] topK(int[] nums, int k) {
    // Java PriorityQueue 默认小顶堆
    PriorityQueue<Integer> minHeap = new PriorityQueue<>(k);
    for (int num : nums) {
        if (minHeap.size() < k) {
            minHeap.offer(num);
        } else if (num > minHeap.peek()) {
            minHeap.poll();
            minHeap.offer(num);
        }
    }
    return minHeap.stream().mapToInt(Integer::intValue).toArray();
}''')

A('追问：如果数据量大到无法全部放入内存怎么办？')
ANS('外部排序 + 分治：')
ABL('将数据分成 M 个文件，每个文件能放入内存')
ABL('每个文件用小顶堆找出 TopK → 得到 M 个候选集')
ABL('合并 M 个候选集，再用小顶堆选出最终 TopK')
ABL('实际场景中可以用 MapReduce/Spark 的 takeOrdered(k) 实现分布式 TopK')

SEC2('8.4 更多高频题速查')
TBL(['题目', '核心思路', '时间复杂度'], [
    ['两数之和', 'HashMap 存已遍历值', 'O(N)'],
    ['反转链表', '三指针 prev/curr/next', 'O(N)'],
    ['有效括号', '栈匹配', 'O(N)'],
    ['合并两个有序链表', '双指针', 'O(N+M)'],
    ['最大子数组和', '动态规划 dp[i]=max(dp[i-1]+a[i], a[i])', 'O(N)'],
    ['二分查找', 'left/right/mid 模板', 'O(log N)'],
    ['岛屿数量', 'DFS/BFS 遍历连通分量', 'O(M×N)'],
    ['最长无重复子串', '滑动窗口 + HashSet', 'O(N)'],
    ['合并区间', '排序后合并', 'O(N log N)'],
    ['手写快速排序', 'partition + 递归', 'O(N log N) 平均'],
])


# ════════════════════════════════════════
# SECTION 9: APPENDIX
# ════════════════════════════════════════
SEC1('九、附录：技术栈速查表')

SEC2('9.1 Java 核心速查')
TBL(['概念', '一句话要点', '面试高频追问'], [
    ['JVM 内存模型', '堆（新生代+老年代）、栈、方法区、程序计数器', '堆溢出排查？GC Roots 有哪些？'],
    ['GC 算法', '标记-清除、标记-复制、标记-整理、分代收集', 'G1 vs ZGC？STW 是什么？'],
    ['HashMap', '数组+链表+红黑树，容量 2^n，负载因子 0.75', '为什么线程不安全？ConcurrentHashMap 原理？'],
    ['ConcurrentHashMap', 'JDK8: CAS+synchronized（锁桶头节点）', '与 Hashtable 的区别？size() 如何计算？'],
    ['线程池', '核心参数：core/max/keepAlive/workQueue/handler', '拒绝策略有哪些？如何合理配置？'],
    ['volatile', '保证可见性+有序性，不保证原子性', 'DCL 单例为什么要用 volatile？'],
    ['synchronized', '偏向锁→轻量锁→重量锁（锁升级）', '与 ReentrantLock 区别？'],
])

SEC2('9.2 Spring Boot 速查')
TBL(['概念', '一句话要点', '面试高频追问'], [
    ['IoC', '控制反转，对象生命周期由容器管理', 'Bean 生命周期？循环依赖？'],
    ['AOP', '动态代理（JDK/CGLIB）实现横切关注点', '事务失效场景？'],
    ['自动配置', '@EnableAutoConfiguration + spring.factories', '如何自定义 Starter？'],
    ['事务', '@Transactional 基于 AOP 代理', '传播行为？事务失效原因？'],
    ['Spring Boot 3', 'Jakarta EE、Virtual Threads、GraalVM', '从 2 升级到 3 的注意事项？'],
])

SEC2('9.3 数据库速查')
TBL(['概念', '一句话要点', '面试高频追问'], [
    ['B+ 树索引', '有序、多路平衡、叶子节点链表', '为什么不用 B 树/红黑树/Hash？'],
    ['MVCC', 'Read View + Undo Log 实现无锁读', 'RR vs RC 的区别？幻读？'],
    ['事务隔离级别', 'RC/RR/Serializable，InnoDB 默认 RR', '各级别分别解决什么问题？'],
    ['慢查询优化', 'EXPLAIN → 索引 → SQL 改写 → 架构优化', 'type 列各值含义？'],
    ['PostgreSQL 特性', 'MVCC、JSONB、pgvector、Apache AGE', '与 MySQL 的核心区别？'],
])

SEC2('9.4 Redis 速查')
TBL(['概念', '一句话要点', '面试高频追问'], [
    ['五种基本类型', 'String/List/Hash/Set/ZSet', '各类型底层编码？'],
    ['缓存穿透', '查询不存在的数据，请求直达 DB', '布隆过滤器/空值缓存'],
    ['缓存雪崩', '大量 key 同时过期', '随机过期时间/多级缓存'],
    ['分布式锁', 'SET key value NX EX', 'Redisson 看门狗？红锁？'],
    ['持久化', 'RDB 快照 + AOF 日志', '混合持久化？数据丢失风险？'],
])

SEC2('9.5 Kafka 速查')
TBL(['概念', '一句话要点', '面试高频追问'], [
    ['消息模型', '发布订阅 + 消费者组', '如何保证顺序？'],
    ['ISR 机制', 'In-Sync Replicas 同步副本集合', 'acks=all 可靠性？'],
    ['Rebalance', '消费者组内分区重新分配', '如何减少 Rebalance？'],
    ['Exactly-Once', '幂等 Producer + 事务', '三种语义区别？'],
    ['高吞吐', '顺序写、零拷贝、批量压缩', '为什么比 RabbitMQ 快？'],
])


# ════════════════════════════════════════
# SECTION 10: 反问面试官
# ════════════════════════════════════════
SEC1('十、反问面试官')
P('面试最后通常会问"你有什么想问的？"，以下是按轮次推荐的问题。好的反问能展示你的思考深度和对岗位的诚意。', sz=10.5, sa=4)

SEC2('10.1 技术面（1-2 面）')
Q('推荐问以下问题，每个都附上为什么要问：')
ANS('1. "团队目前的技术栈是什么？近期有什么技术升级计划？"')
ABL('目的：了解技术方向是否与你的专长匹配，同时展示你关心技术演进')
ANS('2. "这个岗位日常工作中最大的技术挑战是什么？"')
ABL('目的：了解真实工作内容，同时可以顺势补充你相关的经验')
ANS('3. "团队的 Code Review 流程是怎样的？有技术分享机制吗？"')
ABL('目的：了解团队工程文化，展示你重视代码质量和学习氛围')
ANS('4. "项目的代码量和团队规模大概是怎样的？"')
ABL('目的：评估工作强度和协作模式')

SEC2('10.2 主管面')
ANS('1. "团队未来 6-12 个月的核心目标是什么？这个岗位在其中扮演什么角色？"')
ABL('目的：了解战略方向和你的角色定位')
ANS('2. "您觉得在这个岗位上做得出色的人，通常具备什么特质？"')
ABL('目的：了解领导的期望，后续可以针对性地展示匹配度')
ANS('3. "团队的人员流动率如何？最近加入的同事适应得怎么样？"')
ABL('目的：侧面了解团队氛围和稳定性')

SEC2('10.3 HR 面')
ANS('1. "这个岗位的职级和晋升路径是怎样的？"')
ABL('目的：了解职业发展空间')
ANS('2. "团队的绩效考核周期和方式是怎样的？"')
ABL('目的：了解激励机制')
ANS('3. "入职后有新人培训或 Buddy 制度吗？"')
ABL('目的：了解入职支持体系')

SEC2('10.4 不建议问的问题')
ABL('薪资福利（HR 之前不要主动问，显得只关心钱）')
ABL('加班频率（可以侧面了解工作强度，但不要直接问"加班多吗"）')
ABL('百度能查到的基础信息（如公司做什么业务）')
ABL('过于私密的问题（面试官个人背景等）')


# ── Save ──
doc.save(OUTPUT)
n_para = len(doc.paragraphs)
n_tbl = len(doc.tables)
print(f'Part 4 done\nSaved {n_para} paragraphs, {n_tbl} tables')
