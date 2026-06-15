"""Interview prep generator - Part 1: Framework + helpers + Java Core"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = r'C:\Users\47583\projects\hisi_dev_tool v5.0\沈江春_面试准备手册_完整版.docx'
doc = Document()

# --- page setup ---
for sec in doc.sections:
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

# --- styles ---
sn = doc.styles['Normal']; sn.font.name = 'Calibri'; sn.font.size = Pt(10.5)
sn.paragraph_format.space_after = Pt(4); sn.paragraph_format.line_spacing = 1.15
for lv, (sz, clr) in enumerate([(22,'1F4E79'),(16,'2E75B6'),(13,'404040'),(11,'404040')], 1):
    hs = doc.styles[f'Heading {lv}']; hs.font.size = Pt(sz); hs.font.bold = True
    hs.font.color.rgb = RGBColor.from_string(clr)
    hs.paragraph_format.space_before = Pt(12 if lv<=2 else 8); hs.paragraph_format.space_after = Pt(6)

# --- helpers ---
def P(text, bold=False, sz=None, clr=None, align=None, sa=None, sn=None):
    p = doc.add_paragraph(style=sn) if sn else doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if sz: r.font.size = Pt(sz)
    if clr: r.font.color.rgb = RGBColor.from_string(clr)
    if align: p.alignment = align
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    return p

def B(text, lv=0):
    p = doc.add_paragraph(style='List Bullet'); p.clear()
    r = p.add_run(text); r.font.size = Pt(10)
    if lv: p.paragraph_format.left_indent = Cm(1.5*lv)
    return p

def TBL(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        c.paragraphs[0].add_run(h).bold = True; c.paragraphs[0].runs[0].font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            c.paragraphs[0].add_run(v).font.size = Pt(10)
    doc.add_paragraph()

def CODE(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    shd = r._element.get_or_add_rPr().makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):'F5F5F5'})
    r._element.get_or_add_rPr().append(shd)

def A(title):
    """追问题目 + 完整答案"""
    P(title, bold=True, sz=11, clr='2E75B6', sa=2)

def ANS(text):
    """答案正文"""
    P(text, sz=10, sa=6)

def ABL(text):
    """答案中的要点"""
    B(text)

def Q(text):
    """问题标题"""
    P(text, bold=True, sz=11, sa=4)

def SEC1(text): doc.add_heading(text, level=1)
def SEC2(text): doc.add_heading(text, level=2)
def SEC3(text): doc.add_heading(text, level=3)

# ============================================================
# TITLE
# ============================================================
P('沈江春 — Java 后端面试全量准备手册（完整版）', bold=True, sz=22, clr='1F4E79', align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
P('所有追问均附完整详尽答案，可直接用于面试准备', sz=11, clr='666666', align=WD_ALIGN_PARAGRAPH.CENTER, sa=16)

# ============================================================
# PART 1: 技术八股文
# ============================================================
SEC1('第一部分：技术八股文（含完整答案）')

# ==================== 一、Java 核心 ====================
SEC2('一、Java 核心')
SEC3('1.1 JVM 内存模型')

Q('Q：JVM 运行时数据区有哪些？各自作用？')
TBL(['区域','线程共享','作用'],[
    ['堆（Heap）','共享','对象实例、数组，GC 主战场'],
    ['方法区/元空间（Metaspace）','共享','类信息、常量池、静态变量（JDK 8+移至本地内存）'],
    ['虚拟机栈（VM Stack）','私有','方法调用栈帧（局部变量表、操作数栈、动态链接、返回地址）'],
    ['本地方法栈（Native Stack）','私有','Native 方法调用'],
    ['程序计数器（PC Register）','私有','当前执行字节码行号'],
])

A('追问 1：堆内存如何分代？为什么分代？')
ANS('JVM 堆内存分为年轻代（Young Generation）和老年代（Old Generation）。年轻代又进一步分为 1 个 Eden 区和 2 个 Survivor 区（S0、S1），默认比例 8:1:1。')
ANS('分代的原因基于「分代假说」（Generational Hypothesis）：')
ABL('弱分代假说：绝大多数对象都是朝生夕灭的（即很快变成垃圾）。')
ABL('强分代假说：熬过越多次 GC 的对象越难以消亡。')
ANS('基于这两个假说，JVM 将堆分为年轻代和老年代，对不同代采用不同的 GC 策略：年轻代使用复制算法（因为大部分对象会被回收，只需复制少量存活对象），老年代使用标记-清除或标记-整理算法（因为存活对象多，复制成本太高）。这种分代设计使得 GC 整体效率大幅提升。')
ANS('对象从年轻代晋升到老年代的条件：(1) 对象在 Survivor 区熬过 MaxTenuringThreshold 次 Minor GC（默认 15）；(2) Survivor 区中相同年龄所有对象大小之和超过 Survivor 空间一半，年龄 >= 该年龄的对象直接进入老年代；(3) 大对象直接进入老年代（-XX:PretenureSizeThreshold）。')

A('追问 2：Metaspace 与 PermGen 的区别？')
ANS('PermGen（永久代）是 JDK 7 及之前版本中方法区的实现，位于 JVM 堆内存中，大小通过 -XX:MaxPermSize 固定上限。它的主要问题：(1) 大小难以预估，容易 OOM:PermGen space；(2) GC 效率低；(3) 字符串常量池也在 PermGen 中，容易导致溢出。')
ANS('JDK 8 开始移除 PermGen，改用 Metaspace（元空间），核心区别：')
ABL('存储位置：PermGen 在 JVM 堆内，Metaspace 在本地内存（Native Memory）中，不受 -Xmx 限制。')
ABL('大小：Metaspace 默认无上限（受系统物理内存限制），可通过 -XX:MaxMetaspaceSize 设置上限。')
ABL('字符串常量池：JDK 7 就已经从 PermGen 移到了堆中，JDK 8 继续保持。')
ABL('GC：Metaspace 的类卸载更高效，在 Full GC 时触发回收。')
ANS('JDK 8 后 OOM 的表现从 java.lang.OutOfMemoryError: PermGen space 变为 java.lang.OutOfMemoryError: Metaspace。常见触发原因：大量动态生成类（CGLIB 代理、Groovy 脚本、反射等），需要通过 -XX:MaxMetaspaceSize 限制或排查类加载泄漏。')

A('追问 3：栈溢出 StackOverflowError 的场景？')
ANS('StackOverflowError 发生在线程请求的栈深度超过虚拟机所允许的最大深度时。常见场景：')
ABL('递归调用没有正确的终止条件，或终止条件在某些输入下无法触发。')
ABL('方法调用链过长，例如深层嵌套的 JSON/XML 解析。')
ABL('局部变量过多导致栈帧过大（每个栈帧包含局部变量表，变量越多栈帧越大，同样栈空间能容纳的栈帧越少）。')
ANS('排查方法：(1) 检查异常堆栈，定位递归入口；(2) 检查递归终止条件是否正确；(3) 考虑改为迭代实现或使用尾递归优化（Java 编译器/JVM 不保证尾递归优化，但 Scala/Kotlin 支持）。')
ANS('与之相关的另一个错误是 OutOfMemoryError: unable to create new native thread，发生在创建线程数超过系统限制时（每个线程需要分配独立的栈空间，默认 1MB，可通过 -Xss 调整）。')

Q('Q：常用的垃圾回收器有哪些？各自的适用场景？')
TBL(['回收器','特点','适用场景'],[
    ['Serial','单线程，Stop-The-World','客户端模式、小堆'],
    ['Parallel (Throughput)','多线程并行，高吞吐','后台批处理'],
    ['CMS','并发标记清除，低延迟（已废弃）','低延迟需求（JDK 14 移除）'],
    ['G1','Region 分区，可预测停顿','JDK 9+ 默认，大堆（6G+）'],
    ['ZGC','几乎零停顿（<10ms）','超大堆（TB 级），JDK 15+ 生产就绪'],
    ['Shenandoah','低延迟，与 ZGC 类似','RedHat 主导，OpenJDK 支持'],
])

A('追问 1：G1 的 Mixed GC 触发条件？')
ANS('G1 的 Mixed GC 是指同时回收年轻代和部分老年代 Region 的 GC。触发条件：当整个堆的占用率达到 InitiatingHeapOccupancyPercent（IHOP，默认 45%）时，G1 开始并发标记周期。并发标记完成后，下一次 GC 就是 Mixed GC，会选择回收价值最高的老年代 Region（即垃圾最多的 Region，通过 Remembered Set 和 RSet 计算每个 Region 的回收收益）。')
ANS('Mixed GC 不是一次性回收所有老年代 Region，而是分多次进行，每次选择一部分 Region，这样每次 GC 的停顿时间可控。如果 Mixed GC 来不及回收，老年代继续增长到接近满时，会退化为 Full GC（单线程或并行的全堆回收，停顿时间很长）。')
ANS('调优建议：如果频繁触发 Full GC，可以降低 IHOP 值让 Mixed GC 更早触发（-XX:InitiatingHeapOccupancyPercent=35），或者增加堆大小。')

A('追问 2：你项目中用的什么 GC？为什么？')
ANS('在我的项目中使用的是 G1 GC，配合 JDK 21。选择 G1 的原因：(1) JDK 9+ 的默认 GC，无需额外配置；(2) 堆大小在 4-8G 范围，G1 的 Region 分区机制表现优秀；(3) 可以通过 -XX:MaxGCPauseMillis 设置目标停顿时间（如 200ms），G1 会自动调整回收策略。')
ANS('在 JDK 21 中 ZGC 也已经非常成熟（分代 ZGC 在 JDK 21 中正式发布），如果对延迟要求极高（P99 < 10ms），可以切换到 ZGC。但在纳税申报系统这种批处理场景中，吞吐量更重要，G1 更合适。')
ANS('具体配置参数：-Xms4g -Xmx4g -XX:+UseG1GC -XX:MaxGCPauseMillis=200 -XX:G1HeapRegionSize=8m -XX:ParallelGCThreads=8 -XX:ConcGCThreads=2')

A('追问 3：什么是三色标记法？并发标记如何解决漏标问题？')
ANS('三色标记法是 GC 并发标记阶段使用的对象遍历算法，将所有对象分为三种颜色：')
ABL('白色：尚未被扫描到的对象。GC 结束后仍为白色的对象将被回收。')
ABL('灰色：对象本身已被扫描，但它引用的其他对象还未全部扫描。')
ABL('黑色：对象及其所有引用都已被扫描完毕。')
ANS('标记过程：初始时所有对象为白色，GC Roots 直接引用的对象标为灰色；然后从灰色集合中取出对象，将其引用的对象标为灰色，自身标为黑色；重复直到灰色集合为空。')
ANS('并发标记的问题：用户线程和 GC 线程同时运行，用户线程可能修改对象引用关系，导致两种错误：')
ABL('漏标（少回收）：一个白色对象被黑色对象新引用，且所有灰色对象到它的旧引用都已删除。这会导致本该存活的对象被错误回收。')
ABL('多标（多回收）：一个黑色对象的引用被删除，但已经被标为黑色。这只会导致浮动垃圾，下次 GC 再回收即可。')
ANS('漏标的解决方案——CMS 使用「增量更新」（Incremental Update）：当黑色对象新增对白色对象的引用时，将这个黑色对象重新标为灰色（写屏障实现）。G1 使用「原始快照」（SATB, Snapshot At The Beginning）：在引用删除时，记录被删除引用指向的白色对象（原始快照），在重新标记阶段重新扫描这些对象。两种方式各有优劣：增量更新可能增加重新标记的工作量，SATB 可能产生更多浮动垃圾。')

Q('Q：如何排查和解决内存泄漏？')
P('排查步骤：', bold=True, sz=10)
ABL('jmap -heap <pid> 查看堆使用概况，观察 Old Gen 是否持续增长。')
ABL('jstat -gcutil <pid> 1000 观察 GC 频率：如果 Full GC 后 Old Gen 占比仍然很高，说明有泄漏。')
ABL('jmap -dump:format=b,file=heap.hprof <pid> 导出堆转储（注意：dump 会触发 Full GC 并暂停应用，生产环境用 -XX:+HeapDumpOnOutOfMemoryError 自动 dump）。')
ABL('MAT（Memory Analyzer Tool）分析 Dominator Tree：找到占用内存最大的对象，展开引用链，找到 GC Root 到泄漏对象的最短路径。')
ABL('定位到具体代码后修复：常见泄漏模式包括集合未清理、ThreadLocal 未 remove、监听器/回调未注销、内部类持有外部类引用、连接/流未关闭等。')
ANS('补充工具：(1) Arthas（阿里开源）可以在线诊断，无需 dump：heapdump / memory 命令；(2) VisualVM 实时监控；(3) JProfiler 支持内存快照对比（找出两次快照之间新增的对象）。')
ANS('生产环境最佳实践：配置 -XX:+HeapDumpOnOutOfMemoryError -XX:HeapDumpPath=/path/to/dumps，当 OOM 时自动 dump；配合 Prometheus + Grafana 监控 Old Gen 使用趋势。')

A('项目关联：简历提到「排查并修复内存泄漏」——准备具体案例')
ANS('建议按 STAR 法则准备：')
ABL('Situation：线上纳税申报服务在月底报税高峰期，运行 2-3 天后响应时间从 200ms 上升到 2s+，伴随频繁 Full GC。')
ABL('Task：在不影响线上业务的情况下，定位并修复内存泄漏。')
ABL('Action：(1) 通过 jstat 观察到 Full GC 后 Old Gen 占用从 60% 逐步上升到 90%；(2) 谨慎地在低峰期执行 jmap dump；(3) MAT 分析 Dominator Tree，发现一个 ConcurrentHashMap 持续增长，key 是租户 ID + 报表类型；(4) 定位代码发现报表计算模块缓存了中间计算结果用于后续查询，但没有设置过期和大小上限。')
ABL('Result：引入 Caffeine 本地缓存替代手动 HashMap，设置 maximumSize=1000 和 expireAfterAccess=30min。上线后 Old Gen 使用率稳定在 40-60%，Full GC 频率恢复正常。')

SEC3('1.2 Java 集合')

Q('Q：HashMap 底层原理？JDK 7 vs JDK 8 的区别？')
TBL(['维度','JDK 7','JDK 8'],[
    ['数据结构','数组 + 链表','数组 + 链表 + 红黑树'],
    ['插入方式','头插法（并发下死循环）','尾插法'],
    ['扩容','先扩容再插入','先插入再扩容'],
    ['树化条件','无','链表长度 >= 8 且数组长度 >= 64'],
    ['线程安全','否（死循环风险）','否（但无死循环）'],
])

ANS('底层原理：HashMap 是基于哈希表的 Map 实现。通过 key 的 hashCode() 计算数组下标（hash & (n-1)，n 为数组长度），将键值对存储在对应的桶（bucket）中。当发生哈希冲突时，使用链表（JDK 7/8）或红黑树（JDK 8）存储冲突元素。默认初始容量 16，负载因子 0.75，当元素数量超过 capacity * loadFactor 时触发扩容（容量翻倍，重新哈希）。')

A('追问 1：为什么阈值是 8？')
ANS('红黑树化的阈值选择 8 是基于泊松分布的统计分析。HashMap 中每个桶的元素数量服从泊松分布（假设哈希函数均匀），在负载因子 0.75 的情况下，一个桶中恰好有 k 个元素的概率：')
ABL('k=0: 0.60653066')
ABL('k=1: 0.30326533')
ABL('k=2: 0.07581633')
ABL('k=3: 0.01263606')
ABL('k=8: 0.00000006（千万分之六）')
ANS('也就是说，链表长度达到 8 的概率极低（约千万分之六），正常使用几乎不会触发树化。选择 8 是在时间和空间开销之间的平衡：红黑树虽然查找 O(log n)，但节点占用空间更大（TreeNode 是 Node 的 2 倍），且维护树结构有额外开销。只有在极端哈希冲突（可能是恶意攻击或哈希函数设计缺陷）时才需要树化。')
ANS('退化条件：当红黑树节点数 <= 6 时退化为链表（UNTREEIFY_THRESHOLD = 6），选择 6 而不是 8 是为了避免在 7 和 8 之间频繁转换（hysteresis）。')

A('追问 2：ConcurrentHashMap 如何保证线程安全？')
ANS('JDK 7 的 ConcurrentHashMap 使用分段锁（Segment），将整个 Map 分成 16 个 Segment（默认），每个 Segment 继承 ReentrantLock，不同 Segment 的操作互不影响，最大并发度为 16。')
ANS('JDK 8 的 ConcurrentHashMap 进行了重大改进，放弃了 Segment，改为 Node 数组 + 链表/红黑树（与 HashMap 结构一致），线程安全通过以下机制保证：')
ABL('CAS 操作：初始化数组和空桶写入使用 CAS（Unsafe.compareAndSwapInt/Reference），无锁并发。')
ABL('synchronized 锁单个桶：当桶中已有元素时（哈希冲突），对链表/红黑树的头节点加 synchronized 锁，锁粒度从 JDK 7 的 Segment（16 个）细化到桶级别（理论上等于数组长度个）。')
ABL('volatile 读：Node 的 val 和 next 字段用 volatile 修饰，保证可见性。')
ABL('size() 使用 CounterCell[] 分散计数（类似 LongAdder），避免 CAS 热点。')
ANS('注意：ConcurrentHashMap 不允许 key 或 value 为 null（而 HashMap 允许），原因是无法区分「key 不存在返回 null」和「value 本身就是 null」，在并发场景下会产生歧义。')

A('追问 3：HashMap 的 key 可以是 null 吗？ConcurrentHashMap 呢？')
ANS('HashMap 允许 key 为 null。null key 的 hash 值固定为 0，总是放在数组的第一个桶（index=0）。HashMap 的 get 方法对 null key 有特殊处理：先判断 key 是否为 null，如果是则直接遍历 table[0] 查找。')
ANS('ConcurrentHashMap 不允许 key 或 value 为 null，put 时会直接抛出 NullPointerException。原因：在并发环境下，ConcurrentHashMap 的 get(key) 返回 null 有两种可能：(1) key 不存在；(2) key 存在但 value 为 null。调用方无法区分这两种情况，可能导致逻辑错误。例如：')
CODE('''if (map.get(key) == null) {\n    // 是 key 不存在？还是 value 就是 null？\n    // 在并发环境下无法用 containsKey + get 来区分（两步操作不是原子的）\n}''')
ANS('设计哲学：Doug Lea（ConcurrentHashMap 作者）认为并发容器应该尽早暴露错误，而不是隐藏歧义。Hashtable 同样不允许 null key/value。')

SEC3('1.3 并发编程')

Q('Q：线程池核心参数有哪些？拒绝策略有哪些？')
CODE('''ThreadPoolExecutor(
    corePoolSize,      // 核心线程数
    maximumPoolSize,   // 最大线程数
    keepAliveTime,     // 非核心线程空闲存活时间
    unit,              // 时间单位
    workQueue,         // 阻塞队列
    threadFactory,     // 线程工厂
    handler            // 拒绝策略
)''')

P('四种拒绝策略：', bold=True, sz=10)
ABL('AbortPolicy（默认）：直接抛出 RejectedExecutionException，调用方感知到任务被拒绝。')
ABL('CallerRunsPolicy：由提交任务的调用者线程执行该任务。好处是不丢任务，且能自然降低提交速度（调用者线程被占用期间无法提交新任务）。')
ABL('DiscardPolicy：静默丢弃任务，不抛异常。适用于允许丢失的场景（如日志采集）。')
ABL('DiscardOldestPolicy：丢弃队列头部（最早入队）的任务，然后重新提交当前任务。')

A('追问 1：线程池大小如何设置？')
ANS('经典公式来自《Java Concurrency in Practice》（Brian Goetz）：')
ABL('CPU 密集型任务（计算为主，如加密、排序）：线程数 = N_cpu + 1。+1 是为了当某个线程因为偶尔的页缺失或其他原因暂停时，额外的线程能确保 CPU 不空闲。')
ABL('IO 密集型任务（网络/磁盘等待为主，如 HTTP 调用、数据库查询）：线程数 = N_cpu * 2 或更高，具体取决于 IO 等待时间占比。经典公式：N_threads = N_cpu * (1 + W/C)，其中 W 是等待时间，C 是计算时间。')
ANS('实际生产环境中的做法：(1) 先用公式估算初始值；(2) 压测调整：通过 JMeter/Gatling 模拟真实负载，观察 CPU 使用率、线程状态、队列长度、响应时间；(3) 监控指标：CPU 使用率 70-80% 为宜，超过 90% 说明线程过多（上下文切换开销大）。')
ANS('获取 CPU 核数：Runtime.getRuntime().availableProcessors()，注意在 Docker 容器中可能读到宿主机核数，需要 JDK 10+ 或设置 -XX:ActiveProcessorCount。')
ANS('注意：不要使用 Executors.newFixedThreadPool() 或 newCachedThreadPool()，因为它们的 workQueue 是无界的（LinkedBlockingQueue），可能导致 OOM。应该手动创建 ThreadPoolExecutor 并指定有界队列。')

A('追问 2：execute() 和 submit() 的区别？')
ANS('execute() 是 Executor 接口定义的方法，接受 Runnable 参数，无返回值。如果任务抛出异常，会直接在执行线程中抛出 UncaughtExceptionHandler。')
ANS('submit() 是 ExecutorService 接口新增的方法，接受 Runnable 或 Callable 参数，返回 Future 对象。关键区别：')
ABL('返回值：execute 无返回值；submit 返回 Future，可以通过 get() 获取任务结果（Callable）或等待完成（Runnable）。')
ABL('异常处理：submit 内部将任务包装为 FutureTask，如果任务抛出异常，异常会被 FutureTask 捕获并存储，直到调用 get() 时才抛出 ExecutionException。如果不调用 get()，异常会被静默吞掉。这是生产环境中常见的 Bug 来源。')
ANS('最佳实践：如果不需要返回值，优先用 execute()（异常不会被吞）；如果需要返回值或异步结果，用 submit() 并确保调用 get() 处理异常。')

A('追问 3：线程池的五种状态？')
ANS('ThreadPoolExecutor 使用一个 AtomicInteger 变量 ctl 同时存储线程池状态（高 3 位）和工作线程数（低 29 位）：')
ABL('RUNNING（111）：接收新任务，处理队列中的任务。值为 -1 << 29 = -536870912。')
ABL('SHUTDOWN（000）：不接收新任务，但处理队列中的已提交任务。调用 shutdown() 进入此状态。')
ABL('STOP（001）：不接收新任务，不处理队列中的任务，中断正在执行的任务。调用 shutdownNow() 进入此状态。')
ABL('TIDYING（010）：所有任务已终止，workerCount 为 0。即将调用 terminated() 钩子方法。')
ABL('TERMINATED（011）：terminated() 方法执行完毕。')
ANS('状态转换路径：RUNNING -> SHUTDOWN（调用 shutdown()）-> TIDYING（队列和线程池都空了）-> TERMINATED；或者 RUNNING -> STOP（调用 shutdownNow()）-> TIDYING（线程池空了）-> TERMINATED。')

Q('Q：synchronized 和 ReentrantLock 的区别？')
TBL(['维度','synchronized','ReentrantLock'],[
    ['实现','JVM 层面（monitorenter/monitorexit 字节码）','API 层面（基于 AQS - AbstractQueuedSynchronizer）'],
    ['锁释放','自动释放（退出同步块或异常时）','手动释放（必须在 finally 中 unlock()，否则死锁）'],
    ['可中断','不可中断，线程只能一直等待','lockInterruptibly() 支持中断等待'],
    ['公平锁','只能非公平','可选公平/非公平（构造器参数 fair=true/false）'],
    ['条件变量','wait/notify/notifyAll（只有一个等待队列）','Condition（可创建多个等待队列，精确唤醒）'],
    ['锁升级','支持：偏向锁 -> 轻量级锁 -> 重量级锁','不支持'],
])

A('追问 1：AQS 的原理？')
ANS('AQS（AbstractQueuedSynchronizer）是 Java 并发包的核心基础框架，ReentrantLock、Semaphore、CountDownLatch、ReadWriteLock 等都基于它实现。')
ANS('核心组成：')
ABL('volatile int state：同步状态。对于 ReentrantLock，state=0 表示未锁定，state>0 表示锁定（值为重入次数）。')
ABL('CLH 变体双向队列：等待获取锁的线程被封装为 Node 节点，加入 FIFO 队列。每个 Node 有 waitStatus 标识状态（SIGNAL=-1 表示后继节点需要唤醒，CANCELLED=1 表示已取消）。')
ANS('加锁流程（非公平锁）：(1) 线程 A 尝试 CAS 将 state 从 0 设为 1，成功则获取锁，设置 exclusiveOwnerThread = A；(2) 线程 B 再来 CAS 必然失败，被封装为 Node 加入 CLH 队列尾部（CAS + 自旋保证线程安全入队）；(3) 线程 B 在队列中自旋检查前驱节点是否为 head，如果不是则 park（LockSupport.park）阻塞；(4) 线程 A 释放锁（state 设为 0），唤醒队列中 head 的后继节点（LockSupport.unpark），被唤醒的线程 B 尝试获取锁。')
ANS('公平锁的区别：在步骤 (1) 中，公平锁会先检查 CLH 队列中是否有等待更久的线程（hasQueuedPredecessors()），如果有则放弃 CAS，直接入队等待，保证 FIFO。非公平锁则直接尝试 CAS 抢锁，可能会插队。')

A('追问 2：锁升级过程？')
ANS('JDK 6 引入了锁升级机制，根据竞争程度自动升级锁的级别，目的是在无竞争时尽量减少锁的开销：')
ABL('无锁状态：对象刚创建，没有任何线程持有锁。')
ABL('偏向锁（Biased Locking）：当第一个线程访问同步块时，在对象头的 Mark Word 中记录该线程 ID（CAS 操作一次）。后续该线程再次进入同步块时，只需检查 Mark Word 中的线程 ID 是否是自己，无需任何同步操作。适用场景：几乎没有竞争，只有一个线程反复进入。')
ABL('轻量级锁（Lightweight Locking）：当第二个线程尝试获取锁时，偏向锁撤销，升级为轻量级锁。线程在栈帧中创建 Lock Record，CAS 将对象头的 Mark Word 指向 Lock Record。线程通过自旋（CAS 循环）等待获取锁，不进入内核态阻塞。适用场景：轻度竞争，锁持有时间很短（自旋期间对方就能释放）。')
ABL('重量级锁（Heavyweight Locking）：自旋一定次数（JDK 6+ 使用自适应自旋）后仍未获取锁，升级为重量级锁。通过 OS 的 mutex（互斥量）实现，未获取锁的线程进入内核态阻塞（上下文切换开销大）。适用场景：激烈竞争，锁持有时间长。')
ANS('注意：JDK 15 中偏向锁被默认禁用（JEP 374），原因是现代应用中并发程度高，偏向锁的撤销成本（需要 STW 暂停）超过其收益。JDK 18+ 彻底移除偏向锁。')

Q('Q：volatile 的作用？和 synchronized 的区别？')
P('volatile 两个核心作用：', bold=True, sz=10)
ABL('可见性保证：当一个线程写入 volatile 变量时，JMM（Java Memory Model）会将该线程的工作内存中的变量值刷新到主内存；当一个线程读取 volatile 变量时，会从主内存重新读取，而不是使用工作内存中的缓存副本。底层通过内存屏障（Memory Barrier）实现。')
ABL('禁止指令重排序：编译器和处理器可能对指令重排序优化。volatile 写操作前插入 StoreStore 屏障（保证写 volatile 前的普通写已刷新），后插入 StoreLoad 屏障（保证 volatile 写对后续读可见）；volatile 读操作前插入 LoadLoad 屏障，后插入 LoadStore 屏障。')
ANS('volatile 不保证原子性：典型的例子是 i++，它实际上是三步操作（读取 i -> 加 1 -> 写回 i），在多线程环境下可能出现丢失更新。解决方案：AtomicInteger（CAS）、synchronized、LongAdder（高并发计数器）。')
ANS('volatile vs synchronized 关键区别：volatile 只保证可见性和有序性，不保证原子性；synchronized 三者都保证。volatile 不会阻塞线程；synchronized 会阻塞。volatile 只能修饰变量；synchronized 可以修饰方法和代码块。')

A('追问 1：DCL 单例为什么要加 volatile？')
ANS('DCL（Double-Checked Locking）单例模式：')
CODE('''class Singleton {\n    private static volatile Singleton instance;\n    public static Singleton getInstance() {\n        if (instance == null) {           // 第一次检查\n            synchronized (Singleton.class) {\n                if (instance == null) {   // 第二次检查\n                    instance = new Singleton();\n                }\n            }\n        }\n        return instance;\n    }\n}''')
ANS('为什么必须加 volatile？因为 new Singleton() 不是原子操作，实际分为三步：(1) 分配内存空间；(2) 调用构造函数初始化对象；(3) 将 instance 指向分配的内存。JVM 可能对步骤 (2) 和 (3) 重排序为 (1)(3)(2)。')
ANS('如果不加 volatile，线程 A 执行到 instance = new Singleton()，重排序后先执行 (3) 将 instance 指向内存（此时对象尚未初始化），线程 B 在第一次检查时发现 instance != null，直接返回一个尚未完成初始化的对象，访问其字段时可能读到默认值（null/0），导致难以排查的 Bug。')
ANS('volatile 通过禁止重排序解决了这个问题：保证 instance = new Singleton() 的三步严格按照 (1)(2)(3) 顺序执行。')

A('追问 2：happens-before 规则有哪些？')
ANS('happens-before 是 JMM（Java Memory Model）的核心概念，定义了操作之间的可见性保证。如果操作 A happens-before 操作 B，则 A 的结果对 B 可见（注意：不是时间上的先后，是内存可见性的保证）。主要规则：')
ABL('程序次序规则：同一线程内，前面的操作 happens-before 后面的操作。')
ABL('volatile 变量规则：volatile 写 happens-before 后续的 volatile 读。')
ABL('锁规则：锁的释放（unlock）happens-before 后续的加锁（lock）。')
ABL('线程启动规则：Thread.start() happens-before 被启动线程的任何操作。')
ABL('线程终止规则：线程的所有操作 happens-before 其他线程检测到该线程终止（join/isAlive）。')
ABL('传递性规则：如果 A hb B 且 B hb C，则 A hb C。')
ABL('线程中断规则：interrupt() 调用 happens-before 检测到中断（isInterrupted/InterruptedException）。')
ABL('对象终结规则：构造函数结束 happens-before finalize() 方法开始。')

Q('Q：ThreadLocal 原理？内存泄漏问题？')
ANS('ThreadLocal 的实现原理：每个 Thread 对象内部持有一个 ThreadLocalMap（自定义的 HashMap 实现），其中 key 是 ThreadLocal 实例（弱引用），value 是存储的值。当调用 threadLocal.set(value) 时，实际上是获取当前线程的 ThreadLocalMap，以 threadLocal 为 key 存入 value。')
ANS('内存泄漏问题：ThreadLocalMap 的 Entry 继承了 WeakReference<ThreadLocal<?>>，key 是弱引用。当 ThreadLocal 实例没有其他强引用时（比如 ThreadLocal 变量出了作用域），GC 会回收这个 ThreadLocal 对象，此时 Entry 的 key 变为 null，但 value 仍然被 Entry 强引用，无法回收。')
ANS('如果线程是线程池中的线程（复用），ThreadLocalMap 不会随线程结束而销毁，这些 value 会一直堆积，造成内存泄漏。')
ANS('解决方案：使用 ThreadLocal 后必须在 finally 块中调用 remove()：')
CODE('''try {\n    threadLocal.set(someValue);\n    // 业务逻辑\n} finally {\n    threadLocal.remove(); // 清除当前线程的 ThreadLocalMap 中的 Entry\n}''')
ANS('ThreadLocalMap 的自身清理机制：get() 和 set() 方法会启发式地清理 key 为 null 的过期 Entry（stale entry），但这只是减少泄漏风险，不能完全替代 remove()。')

SEC3('1.4 Java 8-21 新特性')

Q('Q：Java 8 的核心新特性？')
ABL('Lambda 表达式：(a, b) -> a + b，函数式编程的基石。配合 @FunctionalInterface 注解定义函数式接口（只有一个抽象方法的接口），如 Runnable、Comparator、自定义接口。')
ABL('Stream API：声明式的集合操作流水线。list.stream().filter(x -> x > 0).map(String::valueOf).collect(Collectors.toList())。支持并行流（parallelStream()），底层使用 ForkJoinPool。中间操作（filter/map/sorted）是惰性求值，终止操作（collect/forEach/reduce）才触发执行。')
ABL('Optional<T>：防空指针利器。Optional.ofNullable(x).map(User::getName).orElse("default")。避免了 if (x != null) 的防御性代码。注意：不要用 Optional 作为方法参数或类字段，只用于返回值。')
ABL('接口默认方法（default method）：接口中可以用 default 关键字定义有实现的方法，解决接口演化问题（给已有接口添加新方法而不破坏实现类）。')
ABL('新的日期时间 API（java.time）：LocalDate（日期）、LocalTime（时间）、LocalDateTime（日期+时间）、ZonedDateTime（带时区）、Duration（时间段）、Period（日期段）。解决了 Date/Calendar 的线程安全和 API 设计问题。')
ABL('方法引用：四种形式——静态方法引用（Integer::parseInt）、实例方法引用（String::length）、对象方法引用（String::toLowerCase）、构造器引用（ArrayList::new）。')

Q('Q：Java 11/17/21 重要特性？')
TBL(['版本','重要特性','实用场景'],[
    ['Java 11','var 局部变量(LTS)、HttpClient(标准化)、String 增强(strip/isBlank/lines/repeat)、Files.readString/writeString','var 简化声明、HttpClient 替代 Apache HttpClient'],
    ['Java 14','switch 表达式(-> 语法)、Records(预览)、NullPointerException 增强消息','switch 返回值、record 简化 POJO'],
    ['Java 15','text blocks(三引号多行字符串)、sealed classes(预览)、隐藏类','SQL/JSON 字符串更可读'],
    ['Java 17','sealed classes(正式)、pattern matching for instanceof(正式)、强封装 JDK 内部 API(LTS)','instanceof 直接类型转换'],
    ['Java 21','virtual threads(正式)、record patterns、sequenced collections、switch pattern matching(LTS)','百万级并发、更强大的模式匹配'],
])

A('追问 1：Virtual Threads 的原理？和平台线程的区别？')
ANS('Virtual Threads（虚拟线程）是 JDK 21 正式发布的 Project Loom 核心特性，目的是解决 Java 高并发场景中线程资源的瓶颈。')
ANS('平台线程（Platform Thread）：传统的 Java 线程，1:1 映射到 OS 内核线程。每个平台线程需要分配独立的栈空间（默认 1MB），OS 内核线程数有上限（通常几千到几万），创建和上下文切换开销大。')
ANS('虚拟线程（Virtual Thread）：由 JVM 调度，N:M 映射到少量平台线程（称为载体线程 carrier thread）。关键特性：')
ABL('极轻量：栈空间按需分配（初始几百字节，动态增长），可创建百万级虚拟线程。')
ABL('JVM 调度：虚拟线程在阻塞操作（IO、Lock、sleep）时自动让出载体线程（unmount），阻塞结束后重新调度到载体线程继续执行。不需要像平台线程那样阻塞 OS 线程。')
ABL('无需 async/await：代码保持同步写法，JVM 在字节码层面自动挂载/卸载虚拟线程。不需要 CompletableFuture 或响应式编程（Reactor/RxJava）的复杂心智负担。')
ANS('使用方式：')
CODE('''// 创建虚拟线程\nThread.startVirtualThread(() -> {\n    // 同步阻塞代码，但底层不阻塞 OS 线程\n    String result = httpClient.send(request);\n    process(result);\n});\n\n// 虚拟线程池（不需要传统线程池）\ntry (var executor = Executors.newVirtualThreadPerTaskExecutor()) {\n    executor.submit(() -> task1());\n    executor.submit(() -> task2());\n}''')
ANS('注意：虚拟线程不适合 CPU 密集型任务（没有 IO 等待就不会让出载体线程）；synchronized 块中的阻塞操作会导致载体线程被 pin（无法卸载），应改用 ReentrantLock。')

A('追问 2：你项目从 JDK 8 升到 21 的主要兼容性问题？')
ANS('主要挑战和解决方案：')
ABL('javax -> jakarta 命名空间迁移：这是最大的工作量。Spring Boot 3 强制使用 Jakarta EE 9+（jakarta.servlet.*、jakarta.persistence.* 等）。所有使用 javax.* 包的代码都需要全局替换。三方库也需要升级到支持 Jakarta 的版本（如 Hibernate 6.x、Tomcat 10+）。')
ABL('反射访问限制：JDK 16+ 默认禁止对 JDK 内部 API 的非法反射访问（--illegal-access=deny）。如果代码通过反射访问 java.lang、java.util 等包的私有字段/方法，需要添加 --add-opens/ --add-exports 参数或重构代码消除反射依赖。')
ABL('Security Manager 废弃（JDK 17 标记废弃，JDK 18 移除相关 API）：如果使用了 SecurityManager 进行权限控制，需要移除相关代码。')
ABL('GC 参数变化：CMS 被移除（JDK 14），如果使用了 -XX:+UseConcMarkSweepGC 需要切换到 G1 或 ZGC。其他废弃参数需要排查。')
ABL('Nashorn JavaScript 引擎移除（JDK 15）：如果有用 ScriptEngine 执行 JS 的逻辑，需要迁移到 GraalVM JS 或 Rhino。')
ABL('String.hashCode() 变化：JDK 12 改进了 String 的哈希算法（从 char 级别改为 byte 级别），如果有序列化了 HashMap<String, ...> 的数据需要考虑兼容性。')

print('Part 1 done, saving checkpoint...')
doc.save(OUTPUT)
print(f'Saved {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables')
